#-----------------------------------------------------------------------------------------------------------------------
# Module:  week_planner.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Week Planner generator for the Year Planner. Generates weekly overview tables using ISO 8601 week numbering.
#   A year has 52 or 53 weeks depending on how the year boundaries fall.
#-----------------------------------------------------------------------------------------------------------------------

from datetime        import date, timedelta
from docx            import Document
from docx.shared     import Cm, Pt, RGBColor, Twips
from docx.enum.text  import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns    import nsdecls, qn
from docx.oxml       import parse_xml

from src.config import Config
from src.document import (
    add_page_break, get_content_width_twips, compute_table_row_height,
    MINIMIZED_PARAGRAPH_HEIGHT_TWIPS, TWIPS_PER_CM, TWIPS_PER_PT,
    add_config_info_overlay, get_title_row_height_twips, get_header_row_height_twips,
    grayscale_to_hex, grayscale_to_rgb
)
from src.utils.styles import FONT_NAME, COLOR_BLACK

# Global constants.

MONTH_NAMES = [
    "", "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December"
]

# Fixed column widths in cm (Week and Month have fixed content width).

WEEK_COL_WIDTH_CM  = 1.4   # For week numbers (1-53)
MONTH_COL_WIDTH_CM = 4.0  # For month names


#-----------------------------------------------------------------------------------------------------------------------
# Function: _calculate_column_widths
#
# Description:
#
#   Calculate column widths dynamically based on content area width.
#
#   Fixed columns: Week (1.4 cm), Month (4.0 cm). Variable columns: Notes split evenly across remaining space.
#
# Arguments:
#
#   config : Configuration with page settings.
#
# Returns:
#
#   Tuple of (week_dxa, month_dxa, notes1_dxa, notes2_dxa).
#-----------------------------------------------------------------------------------------------------------------------

def _calculate_column_widths ( config: Config ) -> tuple [ int, int, int, int ]:

    # Calculate column widths dynamically based on content area width.

    total_width_twips = get_content_width_twips ( config )

    # Fixed column widths in twips.

    week_col_twips  = int ( WEEK_COL_WIDTH_CM * TWIPS_PER_CM )
    month_col_twips = int ( MONTH_COL_WIDTH_CM * TWIPS_PER_CM )

    # Remaining space for Notes columns (split evenly).

    remaining_twips = total_width_twips - week_col_twips - month_col_twips
    notes_col_twips = remaining_twips // 2

    # Adjust for any rounding (add remainder to first notes column).

    notes1_twips = notes_col_twips + ( remaining_twips % 2 )
    notes2_twips = notes_col_twips

    # Return data to caller.

    return ( week_col_twips, month_col_twips, notes1_twips, notes2_twips )


#-----------------------------------------------------------------------------------------------------------------------
# Function: generate_week_planner
#
# Description:
#
#   Generate the Week Planner section.
#
#   Creates tables showing all ISO 8601 weeks (52 or 53) with week number, month, and notes columns.
#   Double-sided with configurable rows per page. The last table only includes rows for remaining weeks
#   (no empty rows).
#
# Arguments:
#
#   document : The Word document to add the week planner to.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def generate_week_planner ( document: Document, config: Config ) -> None:

    # Generate the Week Planner section.

    year          = config.document.year
    rows_per_page = config.raw.get ( 'week_planner', {} ).get ( 'rows_per_page', 14 )

    # Get all ISO 8601 weeks for the year.

    weeks = _get_year_weeks ( year )

    # Generate pages.

    total_weeks = len ( weeks )
    page_count  = 0

    for start_idx in range ( 0, total_weeks, rows_per_page ):

        # Add minimized page break between Week Planner pages (not before first).
        # The first page's page break comes from main.py.

        if page_count > 0:
            add_page_break ( document, minimize_height = True )

        end_idx    = min ( start_idx + rows_per_page, total_weeks )
        page_weeks = weeks [ start_idx : end_idx ]

        # Determine actual row count for this page.
        # Last page only gets as many rows as there are remaining weeks.

        is_last_page = ( end_idx == total_weeks )
        actual_rows  = len ( page_weeks ) if is_last_page else rows_per_page

        _create_week_planner_page ( document, config, page_weeks, rows_per_page,
                                    actual_rows, year )

        # Add overlay — first page is recto, then alternates.

        is_recto = ( page_count % 2 == 0 )
        add_config_info_overlay ( document, config, is_recto = is_recto )

        page_count += 1


#-----------------------------------------------------------------------------------------------------------------------
# Function: _get_year_weeks
#
# Description:
#
#   Get all ISO 8601 weeks for a year with their month assignments.
#
#   Uses ISO 8601 week numbering where:
#   - Weeks start on Monday.
#   - Week 1 is the week containing the first Thursday of the year
#     (equivalently, the week containing January 4th).
#   - A year has 52 or 53 weeks.
#
#   Note: ISO weeks at year boundaries may include days from adjacent calendar years.
#   For example, ISO week 1 of 2026 starts on Monday, December 29, 2025.
#
# Arguments:
#
#   year : The year to get weeks for.
#
# Returns:
#
#   List of tuples (week_number, month_string, is_first_week_of_month) for each week.
#   is_first_week_of_month is True if the week contains the 1st day of any month.
#-----------------------------------------------------------------------------------------------------------------------

def _get_year_weeks ( year: int ) -> list [ tuple [ int, str, bool ] ]:

    # Get all ISO 8601 weeks for a year with their month assignments.

    weeks = []

    # Find the Monday of ISO week 1 for this year.
    # ISO week 1 always contains January 4th.

    jan4         = date ( year, 1, 4 )
    week1_monday = jan4 - timedelta ( days = jan4.weekday () )

    current_monday = week1_monday

    # Iterate through all weeks of this ISO year.

    while True:

        # Use Thursday to determine ISO year (Thursday is always in the
        # correct ISO year for that week).

        current_thursday      = current_monday + timedelta ( days = 3 )
        iso_year, iso_week, _ = current_thursday.isocalendar ()

        # Stop when we have moved past this ISO year.

        if iso_year > year:
            break

        current_sunday = current_monday + timedelta ( days = 6 )
        month_str      = _get_week_months ( current_monday, current_sunday )

        # A week is the first week of a month if:
        # - Monday is the 1st of a month, OR
        # - The week spans two months (a new month starts during the week).

        is_first_week = ( current_monday.day == 1 ) or ( current_monday.month != current_sunday.month )

        weeks.append ( ( iso_week, month_str, is_first_week ) )

        current_monday += timedelta ( days = 7 )

    # Return data to caller.

    return weeks


#-----------------------------------------------------------------------------------------------------------------------
# Function: _get_week_months
#
# Description:
#
#   Get the month string for a week.
#
#   If the week spans two months, returns "Month / Month".
#
# Arguments:
#
#   monday : The Monday of the week.
#   sunday : The Sunday of the week.
#
# Returns:
#
#   Month string (e.g., "January" or "March / April").
#-----------------------------------------------------------------------------------------------------------------------

def _get_week_months ( monday: date, sunday: date ) -> str:

    # Get the month string for a week.

    start_month = monday.month
    end_month   = sunday.month

    if start_month == end_month:

        # Return data to caller.

        return MONTH_NAMES [ start_month ]

    else:

        # Return data to caller.

        return f"{MONTH_NAMES[start_month]} / {MONTH_NAMES[end_month]}"


#-----------------------------------------------------------------------------------------------------------------------
# Function: _create_week_planner_page
#
# Description:
#
#   Create a single page of the week planner.
#
# Arguments:
#
#   document      : The Word document.
#   config        : Configuration settings.
#   weeks         : List of (week_number, month_string, is_first_week_of_month) tuples for this page.
#   rows_per_page : Number of content rows used for row height calculation
#                   (ensures consistent row heights across all pages).
#   actual_rows   : Actual number of content rows to create in this table.
#                   For the last page, this may be less than rows_per_page.
#   year          : The planner year.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _create_week_planner_page ( document: Document, config: Config,
                                weeks: list [ tuple [ int, str, bool ] ],
                                rows_per_page: int, actual_rows: int,
                                year: int ) -> None:

    # Create a single page of the week planner.

    # Create table: title row + header row + content rows.
    # 4 grid columns: Week, Month, Notes col1, Notes col2.

    total_rows      = 2 + actual_rows  # title + header + content
    table           = document.add_table ( rows = total_rows, cols = 4 )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = False

    # Set fixed table layout.

    _set_table_layout_fixed ( table )

    # Calculate column widths dynamically based on content area.

    week_col, month_col, notes1_col, notes2_col = _calculate_column_widths ( config )

    # Set table grid column widths.

    _set_table_grid ( table, [ week_col, month_col, notes1_col, notes2_col ] )

    # Get row heights from config.

    title_row_height_twips  = get_title_row_height_twips ( config )
    header_row_height_twips = get_header_row_height_twips ( config )

    # Compute content row height dynamically to fill exact vertical content area.
    # Formula: p_v = p_para + r_t + r_h + r_c * n
    # Where p_para accounts for the minimized paragraph created by page break.

    content_row_height_twips = compute_table_row_height (
        config                           = config,
        num_content_rows                 = rows_per_page,
        title_row_height_twips           = title_row_height_twips,
        header_row_height_twips          = header_row_height_twips,
        preceding_paragraph_height_twips = MINIMIZED_PARAGRAPH_HEIGHT_TWIPS
    )

    # Get colors from config.

    title_bg_hex     = grayscale_to_hex ( config.table.title_row.background_grayscale )
    title_font_rgb   = grayscale_to_rgb ( config.table.title_row.font_grayscale )
    title_font_color = RGBColor ( *title_font_rgb )
    title_font_size  = Pt ( config.table.title_row.font_size )

    header_bg_hex     = grayscale_to_hex ( config.table.header_row.background_grayscale )
    header_font_rgb   = grayscale_to_rgb ( config.table.header_row.font_grayscale )
    header_font_color = RGBColor ( *header_font_rgb )
    header_font_size  = Pt ( config.table.header_row.font_size )

    # Note: content_row.font_* settings are for optional supplementary text in
    # normally-blank cells, not for structured data like week numbers and months.
    # Structured data uses black text (RGBColor(0, 0, 0)) directly.

    # First week of month shading color from config.

    first_week_grayscale = config.raw.get ( 'week_planner', {} ).get ( 'first_week_grayscale', 5 )
    first_week_bg_hex    = grayscale_to_hex ( first_week_grayscale )

    # === TITLE ROW ===

    title_row = table.rows [ 0 ]
    _set_row_height ( title_row, title_row_height_twips )

    # Merge cells 0-2 for "Week Planner" (gridSpan=3).

    title_cell = title_row.cells [ 0 ]
    title_cell.merge ( title_row.cells [ 2 ] )

    # Set cell widths.

    _set_cell_width ( title_row.cells [ 0 ], week_col + month_col + notes1_col )
    _set_cell_width ( title_row.cells [ 3 ], notes2_col )

    # Title cell: "Week Planner" — config background, config font, bold, left aligned.

    _set_cell_shading ( title_cell, title_bg_hex )
    _set_cell_vertical_alignment ( title_cell, "center" )
    _add_cell_text ( title_cell, "Week Planner", size = title_font_size, bold = True,
                     color = title_font_color, align = None )  # None = default LEFT

    # Year cell — config background, config font, bold, right aligned.

    year_cell = title_row.cells [ 3 ]
    _set_cell_shading ( year_cell, title_bg_hex )
    _set_cell_vertical_alignment ( year_cell, "center" )
    _add_cell_text ( year_cell, str ( year ), size = title_font_size, bold = True,
                     color = title_font_color, align = WD_ALIGN_PARAGRAPH.RIGHT )

    # === HEADER ROW ===

    header_row = table.rows [ 1 ]
    _set_row_height ( header_row, header_row_height_twips )

    # Merge Notes columns (2-3) (gridSpan=2).

    header_row.cells [ 2 ].merge ( header_row.cells [ 3 ] )

    # Set cell widths and content.

    header_widths = [ week_col, month_col, notes1_col + notes2_col ]
    headers       = [ "Week", "Month", "Notes" ]

    header_cells = [ header_row.cells [ 0 ], header_row.cells [ 1 ], header_row.cells [ 2 ] ]
    for i, ( header_text, width ) in enumerate ( zip ( headers, header_widths ) ):
        cell = header_cells [ i ]
        _set_cell_width ( cell, width )
        _set_cell_shading ( cell, header_bg_hex )
        _set_cell_vertical_alignment ( cell, "center" )
        _add_cell_text ( cell, header_text, size = header_font_size, bold = True,
                         color = header_font_color, align = None )  # None = default LEFT

    # === CONTENT ROWS ===

    for row_idx in range ( actual_rows ):
        row = table.rows [ row_idx + 2 ]
        _set_row_height ( row, content_row_height_twips )

        # Merge Notes columns (2-3) (gridSpan=2).

        row.cells [ 2 ].merge ( row.cells [ 3 ] )

        # Set cell widths.

        _set_cell_width ( row.cells [ 0 ], week_col )
        _set_cell_width ( row.cells [ 1 ], month_col )
        _set_cell_width ( row.cells [ 2 ], notes1_col + notes2_col )

        # Set vertical alignment for all cells.

        _set_cell_vertical_alignment ( row.cells [ 0 ], "center" )
        _set_cell_vertical_alignment ( row.cells [ 1 ], "center" )
        _set_cell_vertical_alignment ( row.cells [ 2 ], "center" )

        # Populate row with week data.

        week_num, month_str, is_first_week = weeks [ row_idx ]

        # Apply shading to first week of month rows.

        if is_first_week:
            _set_cell_shading ( row.cells [ 0 ], first_week_bg_hex )
            _set_cell_shading ( row.cells [ 1 ], first_week_bg_hex )
            _set_cell_shading ( row.cells [ 2 ], first_week_bg_hex )

        # Week number — structured data uses black text (not content_row font settings).
        # content_row.font_* is for supplementary text in normally-blank cells.

        _add_cell_text ( row.cells [ 0 ], str ( week_num ), size = None,
                         color = RGBColor ( 0, 0, 0 ), align = None )

        # Month — structured data uses black text.

        _add_cell_text ( row.cells [ 1 ], month_str, size = None,
                         color = RGBColor ( 0, 0, 0 ), align = None )

        # Notes column left empty.

    # Apply borders.

    _set_table_borders ( table, config )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_layout_fixed
#
# Description:
#
#   Set the table layout to fixed so columns do not auto-resize.
#
# Arguments:
#
#   table : The table to set layout on.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_layout_fixed ( table ) -> None:

    # Set the table layout to fixed.

    tbl    = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml (
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert ( 0, tbl_pr )

    tbl_layout = parse_xml (
        '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="fixed"/>'
    )
    tbl_pr.append ( tbl_layout )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_grid
#
# Description:
#
#   Set the table grid column widths.
#
# Arguments:
#
#   table      : The table to set the grid on.
#   col_widths : List of column widths in dxa (twips).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_grid ( table, col_widths: list [ int ] ) -> None:

    # Set the table grid column widths.

    tbl = table._tbl

    # Remove the existing grid if any.

    existing_grid = tbl.find ( qn ( 'w:tblGrid' ) )
    if existing_grid is not None:
        tbl.remove ( existing_grid )

    # Create a new grid.

    grid_xml = '<w:tblGrid xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for width in col_widths:
        grid_xml += f'<w:gridCol w:w="{width}"/>'
    grid_xml += '</w:tblGrid>'

    tbl_grid = parse_xml ( grid_xml )

    # Insert after tblPr.

    tbl_pr = tbl.tblPr
    if tbl_pr is not None:
        tbl_pr.addnext ( tbl_grid )
    else:
        tbl.insert ( 0, tbl_grid )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_row_height
#
# Description:
#
#   Set the row height in twips.
#
# Arguments:
#
#   row          : The table row.
#   height_twips : Height in twips (dxa).
#   exact        : If True, row height is exactly the specified value (won't grow).
#                  If False, row height is at least the specified value (can grow).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_row_height ( row, height_twips: int, exact: bool = True ) -> None:

    # Set the row height in twips.

    tr    = row._tr
    tr_pr = tr.get_or_add_trPr ()

    # Remove the existing height if any.

    existing_height = tr_pr.find ( qn ( 'w:trHeight' ) )
    if existing_height is not None:
        tr_pr.remove ( existing_height )

    h_rule = "exact" if exact else "atLeast"
    tr_height = parse_xml (
        f'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{height_twips}" w:hRule="{h_rule}"/>'
    )
    tr_pr.append ( tr_height )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_width
#
# Description:
#
#   Set the cell width in dxa (twips).
#
# Arguments:
#
#   cell      : The table cell.
#   width_dxa : Width in dxa.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_width ( cell, width_dxa: int ) -> None:

    # Set the cell width in dxa (twips).

    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr ()

    # Remove the existing width if any.

    existing_width = tc_pr.find ( qn ( 'w:tcW' ) )
    if existing_width is not None:
        tc_pr.remove ( existing_width )

    tc_w = parse_xml (
        f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:w="{width_dxa}" w:type="dxa"/>'
    )
    tc_pr.insert ( 0, tc_w )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_vertical_alignment
#
# Description:
#
#   Set the vertical alignment of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   alignment : Alignment value ("top", "center", "bottom").
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_vertical_alignment ( cell, alignment: str ) -> None:

    # Set the vertical alignment of a cell.

    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr ()

    # Remove the existing vAlign if any.

    existing_valign = tc_pr.find ( qn ( 'w:vAlign' ) )
    if existing_valign is not None:
        tc_pr.remove ( existing_valign )

    v_align = parse_xml (
        f'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{alignment}"/>'
    )
    tc_pr.append ( v_align )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_shading
#
# Description:
#
#   Set the background shading color of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   color_hex : Hex color string (e.g., "000000" for black).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_shading ( cell, color_hex: str ) -> None:

    # Set the background shading color of a cell.

    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr ()
    shd = parse_xml (
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tc_pr.append ( shd )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_cell_text
#
# Description:
#
#   Add formatted text to a table cell.
#
# Arguments:
#
#   cell  : The table cell.
#   text  : The text to add.
#   size  : Font size (None to inherit from the document default).
#   bold  : Whether to make the text bold.
#   color : Font color (RGBColor).
#   align : Paragraph alignment (None for default LEFT).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_cell_text ( cell, text: str, size = None, bold: bool = False,
                     color = COLOR_BLACK, align = None ) -> None:

    # Add formatted text to a table cell.

    para = cell.paragraphs [ 0 ]
    if align is not None:
        para.alignment = align
    run           = para.add_run ( text )
    run.font.name = FONT_NAME
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = bold
    run.font.color.rgb = color


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_borders
#
# Description:
#
#   Set table borders using config settings.
#
# Arguments:
#
#   table  : The table to set borders on.
#   config : Configuration with border settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_borders ( table, config: Config ) -> None:

    # Set table borders using config settings.

    # Get border settings from config.

    border_color_hex = grayscale_to_hex ( config.table.border.grayscale )
    border_size      = int ( config.table.border.thickness * 8 )  # Eighths of a point

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml (
        r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    tbl_borders = parse_xml (
        f'''<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:left w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:right w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideH w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideV w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
        </w:tblBorders>'''
    )

    tbl_pr.append ( tbl_borders )
    if tbl.tblPr is None:
        tbl.insert ( 0, tbl_pr )
