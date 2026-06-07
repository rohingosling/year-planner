#-----------------------------------------------------------------------------------------------------------------------
# Module:  instructions.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Instructions page generator for Year Planner. Generates a single page with the instructions image and a title
#   overlay.
#-----------------------------------------------------------------------------------------------------------------------

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
from docx.shared import Cm, Pt

from src.config import Config
from src.document import add_page_break, add_config_info_overlay, get_content_width, get_content_height
from src.paths import resource_path


# Global constants.

IMAGE_DIR  = resource_path("assets", "images")  # Bundled image directory (source root or frozen bundle).
EMU_PER_CM = 914400 / 2.54                      # EMU conversion constant (914400 EMUs per inch, 1 inch = 2.54 cm).


#-----------------------------------------------------------------------------------------------------------------------
# Function: generate_instructions_page
#
# Description:
#
#   Generate the instructions page.
#
#   Single page with instructions image and title overlay. Starts on odd physical page (recto).
#
# Arguments:
#
#   document : The Word document to add the instructions to.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def generate_instructions_page(document: Document, config: Config) -> None:

    # Generate the instructions page.

    # Get content area dimensions.

    content_width_cm = get_content_width(config, include_gutter=True)
    content_height_cm = get_content_height(config)

    # Build image path.

    image_path = IMAGE_DIR / "instructions.png"

    # Insert image filling content area.

    document.add_picture(
        str(image_path),
        width=Cm(content_width_cm),
        height=Cm(content_height_cm)
    )

    # Remove spacing from picture paragraph.

    picture_para = document.paragraphs[-1]
    picture_para.paragraph_format.space_before = Pt(0)
    picture_para.paragraph_format.space_after = Pt(0)

    # Add title overlay (floating text box, transparent background).

    _add_title_overlay(picture_para, config, "Instructions")

    # Add config info overlay (recto) - anchor to picture paragraph since page is full.

    add_config_info_overlay(document, config, is_recto=True, anchor_paragraph=picture_para)

    # Add blank verso page.
    # Note: The page break paragraph is on the CURRENT page (recto), not the blank page.
    # Do NOT anchor the verso overlay to page_break_para - it would appear on the recto.
    # Instead, let add_config_info_overlay create its own minimal paragraph on the verso.

    add_page_break(document, minimize_height=True)
    add_config_info_overlay(document, config, is_recto=False)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_title_overlay
#
# Description:
#
#   Add a floating title text box overlay on top of the image.
#
#   Creates a DrawingML anchored text box with:
#   - Transparent background (no fill).
#   - No border.
#   - Centered text at top of page.
#
# Arguments:
#
#   paragraph  : The paragraph to anchor the text box to.
#   config     : Configuration with page settings.
#   title_text : The title text to display.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_title_overlay(paragraph, config: Config, title_text: str) -> None:

    # Add a floating title text box overlay on top of the image.

    import random

    # Title dimensions and position.

    title_width_cm = 10.0  # Width of title text box.
    title_height_cm = 1.2  # Height to accommodate 18pt text.

    # Calculate center position.

    page_width_cm = config.page.width
    margin_top_cm = config.page.margin_top
    margin_left_cm = config.page.margin_left
    gutter_cm = config.page.gutter_size

    # X position: center of content area (accounting for gutter on recto).
    # Content starts at margin_left + gutter, width is page - margins - gutter.

    content_start_x = margin_left_cm + gutter_cm
    content_width = page_width_cm - margin_left_cm - config.page.margin_right - gutter_cm
    center_x_cm = content_start_x + (content_width - title_width_cm) / 2

    # Y position: at top margin.

    y_cm = margin_top_cm

    # Convert to EMUs.

    x_emu = int(center_x_cm * EMU_PER_CM)
    y_emu = int(y_cm * EMU_PER_CM)
    width_emu = int(title_width_cm * EMU_PER_CM)
    height_emu = int(title_height_cm * EMU_PER_CM)

    # Font size in half-points (18pt = 36 half-points).

    font_size_half_pt = 36

    # Unique ID for this text box.

    doc_pr_id = random.randint(10000, 99999)

    # Build DrawingML XML - using pattern from document.py with transparent fill.

    textbox_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
        xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
      <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                 simplePos="0" relativeHeight="251660000" behindDoc="0"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{doc_pr_id}" name="TitleOverlay_{doc_pr_id}"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks/></wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBox="1"><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
              <wps:spPr bwMode="auto">
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:noFill/>
                <a:ln><a:noFill/></a:ln>
              </wps:spPr>
              <wps:txbx><w:txbxContent><w:p><w:pPr><w:jc w:val="center"/><w:spacing w:after="0" w:line="432" w:lineRule="exact"/></w:pPr><w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/><w:b/><w:sz w:val="{font_size_half_pt}"/><w:szCs w:val="{font_size_half_pt}"/><w:color w:val="000000"/></w:rPr><w:t>{title_text}</w:t></w:r></w:p></w:txbxContent></wps:txbx>
              <wps:bodyPr rot="0" vert="horz" wrap="square" lIns="91440" tIns="45720" rIns="91440" bIns="45720" anchor="t" anchorCtr="0" upright="1"><a:spAutoFit/></wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>'''

    # Parse and add to paragraph.

    drawing = parse_xml(textbox_xml)

    # Add a run to the paragraph and append the drawing.

    run = paragraph.add_run()
    run.font.size = Pt(1)  # Minimal size for the run itself.
    run._r.append(drawing)
