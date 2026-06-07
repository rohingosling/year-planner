#-----------------------------------------------------------------------------------------------------------------------
# Module:  monthly.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Monthly sections generator for the Year Planner. Generates month cover pages and daily spreads for each month.
#-----------------------------------------------------------------------------------------------------------------------

import calendar
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from src.config import Config
from src.document import (
    add_page_break, add_config_info_overlay, get_content_width_twips,
    get_content_height_twips, TWIPS_PER_CM, TWIPS_PER_PT,
    get_title_row_height_twips, get_header_row_height_twips,
    grayscale_to_hex, grayscale_to_rgb, MINIMIZED_PARAGRAPH_HEIGHT_TWIPS,
    SAFETY_MARGIN_TWIPS
)
from src.utils.styles import FONT_NAME, FONT_SIZE_TITLE, COLOR_BLACK


# Global constants.

# Month names for cover pages.

MONTH_NAMES = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December"
]

# Day names.

DAY_NAMES = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]


#-----------------------------------------------------------------------------------------------------------------------
# Function: generate_monthly_sections
#
# Description:
#
#   Generate all monthly sections (12 months).
#
#   Each month contains:
#   - Month cover page (month name centered, recto)
#   - Blank verso
#   - Daily spread pages (always ends on verso)
#
#   The daily spread function guarantees it ends on verso, so each month starts fresh on recto without needing to
#   track page position across months.
#
# Arguments:
#
#   document : The Word document to add the monthly sections to.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def generate_monthly_sections(document: Document, config: Config) -> None:

    # Generate all monthly sections (12 months).

    year = config.document.year

    for month_num in range(1, 13):
        month_name = MONTH_NAMES[month_num - 1]

        # === MONTH COVER (on recto) ===
        # Each month starts on recto because previous month's daily spread guarantees it ends on verso.
        # First month starts on recto from main.py.

        _generate_month_cover(document, config, month_name)
        add_config_info_overlay(document, config, is_recto=True)

        # === BLANK VERSO (after cover) ===

        add_page_break(document)
        add_config_info_overlay(document, config, is_recto=False)

        # === DAILY SPREAD (starts on recto, guarantees end on verso) ===

        add_page_break(document)
        _generate_daily_spread(document, config, year, month_num)

        # Daily spread guarantees it ends on verso. Add MINIMIZED page break to get to recto for next month's cover.
        # Must be minimized because daily spread pages are full (tables fill them).

        if month_num < 12:
            add_page_break(document, minimize_height=True)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _generate_month_cover
#
# Description:
#
#   Generate a month cover page.
#
#   Layout:
#   - Vertically centered month name (36pt, bold, centered)
#
#   The month name is positioned approximately 1/3 down the page for visual balance.
#
# Arguments:
#
#   document   : The Word document.
#   config     : Configuration with document settings.
#   month_name : Name of the month (e.g., "January").
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _generate_month_cover(document: Document, config: Config, month_name: str) -> None:

    # Generate a month cover page.

    # Add spacing paragraphs for vertical positioning. Position month name approximately 1/3 down the page using
    # ~12 empty lines to push content down.

    for _ in range(12):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Month name (36pt, bold, centered).

    month_para = document.add_paragraph()
    month_run = month_para.add_run(month_name)
    month_run.font.name = FONT_NAME
    month_run.font.size = FONT_SIZE_TITLE
    month_run.font.bold = True
    month_run.font.color.rgb = COLOR_BLACK
    month_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


#-----------------------------------------------------------------------------------------------------------------------
# Function: _generate_daily_spread
#
# Description:
#
#   Generate daily spread pages for a month.
#
#   Layout: 2 day tables per page side, double-sided (4 tables per sheet). Always ends on verso (adds blank verso if
#   needed).
#
#   Page position is calculated mathematically based on number of days, not tracked incrementally (to avoid issues
#   with overlay anchor paragraphs).
#
# Arguments:
#
#   document : The Word document.
#   config   : Configuration with document settings.
#   year     : The year.
#   month    : The month number (1-12).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _generate_daily_spread(document: Document, config: Config, year: int, month: int) -> None:

    # Generate daily spread pages for a month.

    # Get config values.

    daily_config = config.raw.get('daily_spread', {})
    num_content_rows = daily_config.get('rows', 8)
    subject_width_percent = daily_config.get('subject_width_percent', 25)
    table_gap_cm = daily_config.get('table_gap', 0.2)

    # Get all days in the month.

    num_days = calendar.monthrange(year, month)[1]
    days = [date(year, month, day) for day in range(1, num_days + 1)]

    # Calculate number of page sides needed (2 tables per page side). This determines the final page position.

    num_page_sides = (num_days + 1) // 2  # Ceiling division for 2 tables per side

    # Calculate table dimensions.

    title_row_height = get_title_row_height_twips(config)
    header_row_height = get_header_row_height_twips(config)
    total_width = get_content_width_twips(config)

    # Calculate content row height for 2 tables per page.

    content_row_height = _calculate_day_table_row_height(
        config, num_content_rows, title_row_height, header_row_height, table_gap_cm
    )

    # Calculate column widths.

    subject_width = int(total_width * subject_width_percent / 100)
    description_width = total_width - subject_width

    # Generate pages with 2 tables each. Track page side for overlay positioning (1=recto, 2=verso, etc.).

    day_idx = 0
    page_side = 1

    while day_idx < len(days):
        is_recto = (page_side % 2 == 1)
        anchor_para = None  # Will be set to gap paragraph if 2 tables on page

        # First table on this page side.

        _create_day_table(
            document, config, days[day_idx],
            total_width, subject_width, description_width,
            title_row_height, header_row_height, content_row_height,
            num_content_rows
        )
        day_idx += 1

        # Add gap and second table (only if there are more days).

        if day_idx < len(days):
            anchor_para = _add_table_gap(document, table_gap_cm)

            _create_day_table(
                document, config, days[day_idx],
                total_width, subject_width, description_width,
                title_row_height, header_row_height, content_row_height,
                num_content_rows
            )
            day_idx += 1

        # Add page break if more days remain.

        if day_idx < len(days):
            page_break_para = add_page_break(document, minimize_height=True)

            # Use page break paragraph as anchor if no gap (single table page).

            if anchor_para is None:
                anchor_para = page_break_para
            page_side += 1

        # Add overlay anchored to existing paragraph (gap or page break). This avoids creating a new paragraph that
        # would overflow the full page.

        if anchor_para is not None:
            add_config_info_overlay(document, config, is_recto=is_recto,
                                    anchor_paragraph=anchor_para)

    # Handle last page if it had only 1 table (no anchor was available). This happens for 31-day months where day 31
    # is alone on the last page. Single-table pages have room for an anchor (only half the page is used).

    if anchor_para is None:
        last_page_is_recto = (page_side % 2 == 1)

        # Create minimal anchor paragraph - won't overflow since page is half empty.

        anchor_para = document.add_paragraph()
        anchor_para.paragraph_format.space_before = Pt(0)
        anchor_para.paragraph_format.space_after = Pt(0)
        anchor_para.paragraph_format.line_spacing = Pt(1)
        anchor_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run = anchor_para.add_run()
        run.font.size = Pt(1)
        add_config_info_overlay(document, config, is_recto=last_page_is_recto,
                                anchor_paragraph=anchor_para)

    # Determine if we ended on recto or verso based on total page sides. Odd number of page sides = ends on recto,
    # even = ends on verso.

    ends_on_recto = (num_page_sides % 2 == 1)

    # Ensure we end on verso (add blank verso if we ended on recto). Must be minimized because daily spread pages are
    # full (tables fill them).

    if ends_on_recto:
        add_page_break(document, minimize_height=True)

        # Create anchor paragraph ON the blank verso (after the page break). The page break paragraph itself is on the
        # previous page (recto).

        blank_verso_anchor = document.add_paragraph()
        blank_verso_anchor.paragraph_format.space_before = Pt(0)
        blank_verso_anchor.paragraph_format.space_after = Pt(0)
        blank_verso_anchor.paragraph_format.line_spacing = Pt(1)
        blank_verso_anchor.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run = blank_verso_anchor.add_run()
        run.font.size = Pt(1)
        add_config_info_overlay(document, config, is_recto=False,
                                anchor_paragraph=blank_verso_anchor)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _calculate_day_table_row_height
#
# Description:
#
#   Calculate content row height for day tables (2 tables per page).
#
#   Formula:
#   available = page_height - margins - safety - paragraph
#   per_table = (available - gap) / 2
#   content_row = (per_table - title - header) / num_rows
#
# Arguments:
#
#   config            : Configuration with page settings.
#   num_content_rows  : Number of content rows per table.
#   title_row_height  : Title row height in twips.
#   header_row_height : Header row height in twips.
#   table_gap_cm      : Gap between tables in cm.
#
# Returns:
#
#   Content row height in twips.
#-----------------------------------------------------------------------------------------------------------------------

def _calculate_day_table_row_height(
    config: Config,
    num_content_rows: int,
    title_row_height: int,
    header_row_height: int,
    table_gap_cm: float
) -> int:

    # Calculate content row height for day tables (2 tables per page).

    # Total available height.

    available_twips = get_content_height_twips(config)

    # Subtract overhead.

    overhead = MINIMIZED_PARAGRAPH_HEIGHT_TWIPS + SAFETY_MARGIN_TWIPS
    available_twips -= overhead

    # Gap between tables in twips.

    gap_twips = int(table_gap_cm * TWIPS_PER_CM)

    # Height available for 2 tables.

    total_for_tables = available_twips - gap_twips

    # Height per table.

    per_table = total_for_tables // 2

    # Subtract title and header rows.

    content_area = per_table - title_row_height - header_row_height

    # Divide by number of content rows and return data to caller.

    return content_area // num_content_rows


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_table_gap
#
# Description:
#
#   Add a gap paragraph between day tables.
#
#   Uses exact line spacing to create a precise gap height.
#
# Arguments:
#
#   document : The Word document.
#   gap_cm   : Gap size in centimeters.
#
# Returns:
#
#   The gap paragraph (can be used as anchor for overlays).
#-----------------------------------------------------------------------------------------------------------------------

def _add_table_gap(document: Document, gap_cm: float):

    # Add a gap paragraph between day tables.

    gap_para = document.add_paragraph()
    gap_para.paragraph_format.space_before = Pt(0)
    gap_para.paragraph_format.space_after = Pt(0)

    # Convert cm to points (1 cm ≈ 28.35 points).

    gap_pt = gap_cm * 28.35

    # Use exact line spacing to create the gap. The paragraph needs minimal content for line spacing to apply.

    gap_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    gap_para.paragraph_format.line_spacing = Pt(gap_pt)

    # Add a space character so the paragraph has content (empty paragraphs can collapse).

    run = gap_para.add_run(" ")
    run.font.size = Pt(1)  # Minimal font size

    # Return data to caller.

    return gap_para


#-----------------------------------------------------------------------------------------------------------------------
# Function: _create_day_table
#
# Description:
#
#   Create a single day table.
#
#   Structure:
#   - Title row: Day name (left) | Date string (right)
#   - Header row: Subject | Description
#   - Content rows: Empty cells for user input
#
# Arguments:
#
#   document           : The Word document.
#   config             : Configuration settings.
#   day                : The date for this table.
#   total_width        : Total table width in twips.
#   subject_width      : Subject column width in twips.
#   description_width  : Description column width in twips.
#   title_row_height   : Title row height in twips.
#   header_row_height  : Header row height in twips.
#   content_row_height : Content row height in twips.
#   num_content_rows   : Number of content rows.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _create_day_table(
    document: Document,
    config: Config,
    day: date,
    total_width: int,
    subject_width: int,
    description_width: int,
    title_row_height: int,
    header_row_height: int,
    content_row_height: int,
    num_content_rows: int
) -> None:

    # Create a single day table.

    # Create table: title + header + content rows.

    total_rows = 2 + num_content_rows
    table = document.add_table(rows=total_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set fixed table layout.

    _set_table_layout_fixed(table)

    # Set column widths.

    col_widths = [subject_width, description_width]
    _set_table_grid(table, col_widths)

    # Get styling from config.

    title_bg_hex = grayscale_to_hex(config.table.title_row.background_grayscale)
    title_font_rgb = grayscale_to_rgb(config.table.title_row.font_grayscale)
    title_font_color = RGBColor(*title_font_rgb)
    title_font_size = Pt(config.table.title_row.font_size)

    header_bg_hex = grayscale_to_hex(config.table.header_row.background_grayscale)
    header_font_rgb = grayscale_to_rgb(config.table.header_row.font_grayscale)
    header_font_color = RGBColor(*header_font_rgb)
    header_font_size = Pt(config.table.header_row.font_size)

    # === TITLE ROW ===

    title_row = table.rows[0]
    _set_row_height(title_row, title_row_height)

    # Day name (left cell).

    day_name = DAY_NAMES[day.weekday()]
    day_cell = title_row.cells[0]
    _set_cell_width(day_cell, subject_width)
    _set_cell_shading(day_cell, title_bg_hex)
    _set_cell_vertical_alignment(day_cell, "center")
    _add_cell_text(day_cell, day_name, size=title_font_size, bold=True,
                   color=title_font_color, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Date string (right cell).

    date_str = _format_date_string(day)
    date_cell = title_row.cells[1]
    _set_cell_width(date_cell, description_width)
    _set_cell_shading(date_cell, title_bg_hex)
    _set_cell_vertical_alignment(date_cell, "center")
    _add_cell_text(date_cell, date_str, size=title_font_size, bold=True,
                   color=title_font_color, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # === HEADER ROW ===

    header_row = table.rows[1]
    _set_row_height(header_row, header_row_height)

    # Subject header.

    subject_cell = header_row.cells[0]
    _set_cell_width(subject_cell, subject_width)
    _set_cell_shading(subject_cell, header_bg_hex)
    _set_cell_vertical_alignment(subject_cell, "center")
    _add_cell_text(subject_cell, "Subject", size=header_font_size, bold=True,
                   color=header_font_color, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Description header.

    desc_cell = header_row.cells[1]
    _set_cell_width(desc_cell, description_width)
    _set_cell_shading(desc_cell, header_bg_hex)
    _set_cell_vertical_alignment(desc_cell, "center")
    _add_cell_text(desc_cell, "Description", size=header_font_size, bold=True,
                   color=header_font_color, align=WD_ALIGN_PARAGRAPH.LEFT)

    # === CONTENT ROWS ===

    for row_idx in range(num_content_rows):
        row = table.rows[row_idx + 2]
        _set_row_height(row, content_row_height)

        # Subject cell.

        subj_cell = row.cells[0]
        _set_cell_width(subj_cell, subject_width)
        _set_cell_vertical_alignment(subj_cell, "center")

        # Description cell.

        desc_cell = row.cells[1]
        _set_cell_width(desc_cell, description_width)
        _set_cell_vertical_alignment(desc_cell, "center")

    # Apply borders.

    _set_table_borders(table, config)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _format_date_string
#
# Description:
#
#   Format date as "Month Nth,  YYYY-MM-DD,  Week #".
#
#   Two spaces after each comma for visual spacing.
#
# Arguments:
#
#   day : The date to format.
#
# Returns:
#
#   Formatted date string (e.g., "January 1st,  2026-01-01,  Week 1").
#-----------------------------------------------------------------------------------------------------------------------

def _format_date_string(day: date) -> str:

    # Format date as "Month Nth,  YYYY-MM-DD,  Week #".

    month_name = MONTH_NAMES[day.month - 1]
    ordinal = _get_ordinal_suffix(day.day)
    iso_date = day.strftime("%Y-%m-%d")
    week_num = day.isocalendar()[1]

    # Return data to caller.

    return f"{month_name} {day.day}{ordinal},  {iso_date},  Week {week_num}"


#-----------------------------------------------------------------------------------------------------------------------
# Function: _get_ordinal_suffix
#
# Description:
#
#   Get ordinal suffix for a number (st, nd, rd, th).
#
# Arguments:
#
#   n : The number.
#
# Returns:
#
#   Ordinal suffix string.
#-----------------------------------------------------------------------------------------------------------------------

def _get_ordinal_suffix(n: int) -> str:

    # Get ordinal suffix for a number (st, nd, rd, th).

    if 11 <= n <= 13:
        return "th"

    # Return data to caller.

    return {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")


# === Table Helper Functions ===

#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_layout_fixed
#
# Description:
#
#   Set table layout to fixed.
#
# Arguments:
#
#   table : The table to set layout on.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_layout_fixed(table) -> None:

    # Set table layout to fixed.

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tbl_pr)

    tbl_layout = parse_xml(
        '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="fixed"/>'
    )
    tbl_pr.append(tbl_layout)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_grid
#
# Description:
#
#   Set the table grid column widths.
#
# Arguments:
#
#   table      : The table to set the grid on.
#   col_widths : List of column widths in dxa (twips).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_grid(table, col_widths: list[int]) -> None:

    # Set the table grid column widths.

    tbl = table._tbl

    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    grid_xml = '<w:tblGrid xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for width in col_widths:
        grid_xml += f'<w:gridCol w:w="{width}"/>'
    grid_xml += '</w:tblGrid>'

    tbl_grid = parse_xml(grid_xml)

    tbl_pr = tbl.tblPr
    if tbl_pr is not None:
        tbl_pr.addnext(tbl_grid)
    else:
        tbl.insert(0, tbl_grid)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_row_height
#
# Description:
#
#   Set the row height in twips.
#
# Arguments:
#
#   row          : The table row.
#   height_twips : Height in twips (dxa).
#   exact        : If True, row height is exactly the specified value (won't grow).
#                  If False, row height is at least the specified value (can grow).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_row_height(row, height_twips: int, exact: bool = True) -> None:

    # Set the row height in twips.

    tr = row._tr
    tr_pr = tr.get_or_add_trPr()

    existing_height = tr_pr.find(qn('w:trHeight'))
    if existing_height is not None:
        tr_pr.remove(existing_height)

    h_rule = "exact" if exact else "atLeast"
    tr_height = parse_xml(
        f'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{height_twips}" w:hRule="{h_rule}"/>'
    )
    tr_pr.append(tr_height)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_width
#
# Description:
#
#   Set cell width in dxa (twips).
#
# Arguments:
#
#   cell      : The table cell.
#   width_dxa : Width in dxa.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_width(cell, width_dxa: int) -> None:

    # Set cell width in dxa (twips).

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    existing_width = tc_pr.find(qn('w:tcW'))
    if existing_width is not None:
        tc_pr.remove(existing_width)

    tc_w = parse_xml(
        f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:w="{width_dxa}" w:type="dxa"/>'
    )
    tc_pr.insert(0, tc_w)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_vertical_alignment
#
# Description:
#
#   Set vertical alignment of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   alignment : Alignment value ("top", "center", "bottom").
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_vertical_alignment(cell, alignment: str) -> None:

    # Set vertical alignment of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    existing_valign = tc_pr.find(qn('w:vAlign'))
    if existing_valign is not None:
        tc_pr.remove(existing_valign)

    v_align = parse_xml(
        f'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{alignment}"/>'
    )
    tc_pr.append(v_align)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_shading
#
# Description:
#
#   Set the background shading color of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   color_hex : Hex color string (e.g., "000000" for black).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str) -> None:

    # Set the background shading color of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tc_pr.append(shd)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_cell_text
#
# Description:
#
#   Add formatted text to a table cell.
#
# Arguments:
#
#   cell  : The table cell.
#   text  : The text to add.
#   size  : Font size (None to inherit from the document default).
#   bold  : Whether to make the text bold.
#   color : Font color (RGBColor).
#   align : Paragraph alignment (None for default LEFT).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_cell_text(cell, text: str, size=None, bold: bool = False,
                   color=COLOR_BLACK, align=None) -> None:

    # Add formatted text to a table cell.

    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = FONT_NAME
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = bold
    run.font.color.rgb = color


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_borders
#
# Description:
#
#   Set table borders using config settings.
#
# Arguments:
#
#   table  : The table to set borders on.
#   config : Configuration with border settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_borders(table, config: Config) -> None:

    # Set table borders using config settings.

    border_color_hex = grayscale_to_hex(config.table.border.grayscale)
    border_size = int(config.table.border.thickness * 8)

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    tbl_borders = parse_xml(
        f'''<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:left w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:right w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideH w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideV w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
        </w:tblBorders>'''
    )

    tbl_pr.append(tbl_borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)
