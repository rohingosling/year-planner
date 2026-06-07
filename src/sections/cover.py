#-----------------------------------------------------------------------------------------------------------------------
# Module:  cover.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Cover page generator for the Year Planner. Generates the front cover (title, version, year)
#   and inside cover (contact information table).
#-----------------------------------------------------------------------------------------------------------------------

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.config import Config
from src.document import get_content_width, get_content_height, add_page_break, add_config_info_overlay
from src.utils.styles import (
    FONT_NAME, FONT_SIZE_TITLE, FONT_SIZE_SUBTITLE, FONT_SIZE_NORMAL,
    COLOR_BLACK
)
from src.utils.tables import (
    create_table, set_table_borders, remove_table_borders,
    set_cell_vertical_alignment, set_cell_shading
)


#-----------------------------------------------------------------------------------------------------------------------
# Function: generate_cover_page
#
# Description:
#
#   Generate the complete cover page (front and inside).
#
# Arguments:
#
#   document : The Word document to add the cover to.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def generate_cover_page(document: Document, config: Config) -> None:

    # Generate the complete cover page (front and inside).

    _generate_front_cover(document, config)
    add_config_info_overlay(document, config, is_recto=True)  # Page 1 (recto)
    add_page_break(document)
    _generate_inside_cover(document, config)
    add_config_info_overlay(document, config, is_recto=False)  # Page 2 (verso)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _generate_front_cover
#
# Description:
#
#   Generate the front (outside) cover page.
#
#   Layout:
#   - 10 empty lines (top spacing).
#   - Document title (36pt, bold, centered).
#   - Version number (14pt, centered).
#   - 2 empty lines.
#   - Year (36pt, bold, centered).
#
# Arguments:
#
#   document : The Word document.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _generate_front_cover(document: Document, config: Config) -> None:

    # Generate the front (outside) cover page.

    # Add 10 empty paragraphs for top spacing.

    for _ in range(10):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title.

    title_para = document.add_paragraph()
    title_run = title_para.add_run(config.document.title)
    title_run.font.name = FONT_NAME
    title_run.font.size = FONT_SIZE_TITLE
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_BLACK
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Version.

    version_para = document.add_paragraph()
    version_run = version_para.add_run(f"Version {config.document.version}")
    version_run.font.name = FONT_NAME
    version_run.font.size = FONT_SIZE_SUBTITLE
    version_run.font.color.rgb = COLOR_BLACK
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2 empty lines before year.

    for _ in range(2):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Year.

    year_para = document.add_paragraph()
    year_run = year_para.add_run(str(config.document.year))
    year_run.font.name = FONT_NAME
    year_run.font.size = FONT_SIZE_TITLE
    year_run.font.bold = True
    year_run.font.color.rgb = COLOR_BLACK
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


#-----------------------------------------------------------------------------------------------------------------------
# Function: _generate_inside_cover
#
# Description:
#
#   Generate the inside cover page with a contact information table.
#
#   Layout:
#   - "If found, please contact:" header (18pt, bold, centered) at top.
#   - Empty paragraph.
#   - Contact information table (labels right-aligned, cells vertically centered).
#
# Arguments:
#
#   document : The Word document.
#   config   : Configuration with cover settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _generate_inside_cover(document: Document, config: Config) -> None:

    # Generate the inside cover page with a contact information table.

    # Header text (18pt, bold, centered) - starts at top of page.

    header_para = document.add_paragraph()
    header_run = header_para.add_run("If found, please contact:")
    header_run.font.name = FONT_NAME
    header_run.font.size = Pt(18)
    header_run.font.bold = True
    header_run.font.color.rgb = COLOR_BLACK
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Empty paragraph before table.

    document.add_paragraph()

    # Create contact information table.

    contact_fields = config.cover.contact_fields
    table_config = config.cover.contact_table
    num_rows = len(contact_fields)

    table = document.add_table(rows=num_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set column widths from config (in cm).

    label_width = Cm(table_config.label_width)
    value_width = Cm(table_config.value_width)
    row_height = Cm(table_config.row_height)

    # Get label cell shading from contact_table config.

    label_bg_grayscale = config.raw.get('cover', {}).get('contact_table', {}).get('label_grayscale', 5)

    for i, field in enumerate(contact_fields):
        row = table.rows[i]
        row.height = row_height

        # Label cell (right-aligned, vertically centered, shaded).

        label_cell = row.cells[0]
        label_cell.width = label_width
        set_cell_vertical_alignment(label_cell, "center")
        set_cell_shading(label_cell, label_bg_grayscale)
        label_para = label_cell.paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label_run = label_para.add_run(field)
        label_run.font.name = FONT_NAME
        label_run.font.bold = True
        label_run.font.color.rgb = COLOR_BLACK

        # Value cell (empty, vertically centered).

        value_cell = row.cells[1]
        value_cell.width = value_width
        set_cell_vertical_alignment(value_cell, "center")

    # Apply table borders.

    set_table_borders(table, config.table)
