#-----------------------------------------------------------------------------------------------------------------------
# Module:  calendar.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Calendar section generator for the Year Planner. Generates two calendar pages:
#
#   - Recto 1: Current year calendar.
#   - Verso 1: Blank.
#   - Recto 2: Next year calendar.
#   - Verso 2: Blank.
#-----------------------------------------------------------------------------------------------------------------------

import calendar as cal_module
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from src.config import Config
from src.document import (
    add_page_break, get_content_width, TWIPS_PER_CM, TWIPS_PER_PT,
    add_config_info_overlay, get_title_row_height_twips,
    grayscale_to_hex, grayscale_to_rgb,
    get_content_height_twips, MINIMIZED_PARAGRAPH_HEIGHT_TWIPS,
    SAFETY_MARGIN_TWIPS, validate_table_height
)
from src.utils.styles import FONT_NAME, COLOR_BLACK
from src.utils.tables import set_table_borders


#-----------------------------------------------------------------------------------------------------------------------
# Function: generate_calendar_section
#
# Description:
#
#   Generate the calendar section with current year and next year calendars.
#
#   Layout (fixed structure regardless of calendar table height):
#   - Page 1 (recto): Current year calendar.
#   - Page 2 (verso): Blank.
#   - Page 3 (recto): Next year calendar.
#   - Page 4 (verso): Blank.
#
#   The section always ends on verso (page 4), so the next section can start on recto
#   with a single page break from main.py.
#
# Arguments:
#
#   document : The Word document to add calendars to.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def generate_calendar_section(document: Document, config: Config) -> None:

    # Generate the calendar section with current year and next year calendars.

    current_year = config.document.year
    next_year = current_year + 1

    # Get calendar settings from config.

    calendar_config = config.raw.get('calendar', {})
    day_row_height_pt = calendar_config.get('day_row_height', 15.5)
    month_name_gap_pt = calendar_config.get('month_name_gap', 6)

    content_width = get_content_width(config)

    # === Page 1 (recto): Current year calendar ===

    _generate_year_calendar_page(document, current_year, content_width,
                                  day_row_height_pt, month_name_gap_pt, config)

    # Add overlay BEFORE page break so it's definitely on the calendar page.

    add_config_info_overlay(document, config, is_recto=True)

    # Page break to move to Page 2 (blank verso).

    add_page_break(document, minimize_height=True)

    # === Page 2 (verso): Blank ===

    # Cursor is now on Page 2 - overlay creates paragraph here.

    add_config_info_overlay(document, config, is_recto=False)

    # Page break to move to Page 3.

    add_page_break(document, minimize_height=True)

    # === Page 3 (recto): Next year calendar ===

    _generate_year_calendar_page(document, next_year, content_width,
                                  day_row_height_pt, month_name_gap_pt, config)

    # Add overlay BEFORE page break so it's definitely on the calendar page.

    add_config_info_overlay(document, config, is_recto=True)

    # Page break to move to Page 4 (blank verso).

    add_page_break(document, minimize_height=True)

    # === Page 4 (verso): Blank ===

    # Cursor is now on Page 4 - overlay creates paragraph here.
    # Section ends on verso so next section starts on recto.

    add_config_info_overlay(document, config, is_recto=False)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _generate_year_calendar_page
#
# Description:
#
#   Generate a single year calendar page with dynamically computed height.
#
#   The calendar height is calculated to fill the available content area, similar to
#   other page-filling tables in the Year Planner.
#
# Arguments:
#
#   document           : The Word document.
#   year               : The year to generate the calendar for.
#   content_width      : Available width in centimeters.
#   day_row_height_pt  : Height of day rows in points.
#   month_name_gap_pt  : Space between month name and days grid in points.
#   config             : Configuration settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _generate_year_calendar_page(document: Document, year: int,
                                  content_width: float,
                                  day_row_height_pt: int, month_name_gap_pt: int,
                                  config: Config) -> None:

    # Generate a single year calendar page with dynamically computed height.

    # Calculate available height dynamically (no preceding paragraph for first table).
    # Calendar table has 1 title row + 6 month rows.

    title_row_height_twips = get_title_row_height_twips(config)

    # Validate table height — this will warn if the table won't fit.
    # Using 0 for header_row_height since calendar doesn't have a header row.

    validate_table_height(
        config,
        section_name=f"Calendar {year}",
        num_content_rows=6,  # 6 month rows
        title_row_height_twips=title_row_height_twips,
        header_row_height_twips=0,  # No header row in calendar
        preceding_paragraph_height_twips=0  # First element on page
    )

    # Create 2x6 grid of month calendars with title row.

    _create_year_calendar_grid(document, year, content_width,
                               day_row_height_pt, month_name_gap_pt, config)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _create_year_calendar_grid
#
# Description:
#
#   Create a 2x6 grid showing all 12 months of a year with a title row.
#
#   The height is calculated dynamically to fill the available content area.
#
#   Layout (2 columns x 7 rows):
#   - Row 0: Title row (year number, spanning both columns).
#   - Row 1: January, February.
#   - Row 2: March, April.
#   - Row 3: May, June.
#   - Row 4: July, August.
#   - Row 5: September, October.
#   - Row 6: November, December.
#
# Arguments:
#
#   document           : The Word document.
#   year               : The year to generate the calendar for.
#   content_width      : Available width in centimeters.
#   day_row_height_pt  : Height of day rows in points.
#   month_name_gap_pt  : Space between month name and days grid in points.
#   config             : Configuration settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _create_year_calendar_grid(document: Document, year: int,
                                content_width: float,
                                day_row_height_pt: int, month_name_gap_pt: int,
                                config: Config) -> None:

    # Create a 2x6 grid showing all 12 months of a year with a title row.

    # Create outer table (1 title row + 6 month rows x 2 columns).

    outer_table = document.add_table(rows=7, cols=2)
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer_table.autofit = False

    # Set fixed table layout to prevent auto-resizing.

    _set_table_layout_fixed(outer_table)

    # Set default cell margins to 0 for left/right (Word defaults to 0.19cm).

    _set_table_cell_margins(outer_table, left=0, right=0)

    # Calculate cell dimensions.

    cell_width = content_width / 2
    content_width_twips = int(content_width * TWIPS_PER_CM)
    cell_width_twips = content_width_twips // 2

    # Set explicit table grid column widths.

    _set_table_grid(outer_table, [cell_width_twips, cell_width_twips])

    # Get title row height from config.

    title_row_height_twips = get_title_row_height_twips(config)

    # Calculate month row height dynamically to fill the page.
    # Formula: available = content_height - title_row - safety_margin - border_overhead.
    # Then: month_row_height = available / 6 (for 6 month rows).
    #
    # Table borders add overhead: 0.5pt borders at top/bottom = ~20 twips.
    # Plus additional safety for Word's implicit table spacing.

    TABLE_BORDER_OVERHEAD_TWIPS = 80  # Extra margin for table borders and spacing
    available_twips = get_content_height_twips(config)
    fixed_overhead = SAFETY_MARGIN_TWIPS + TABLE_BORDER_OVERHEAD_TWIPS + title_row_height_twips
    remaining_twips = available_twips - fixed_overhead
    month_row_height_twips = int(remaining_twips / 6)

    # Get title row styling from config.

    title_bg_hex = grayscale_to_hex(config.table.title_row.background_grayscale)
    title_font_rgb = grayscale_to_rgb(config.table.title_row.font_grayscale)
    title_font_color = RGBColor(*title_font_rgb)
    title_font_size = Pt(config.table.title_row.font_size)

    # Get header row background for month name shading.

    header_bg_hex = grayscale_to_hex(config.table.header_row.background_grayscale)

    # === TITLE ROW ===

    title_row = outer_table.rows[0]
    _set_row_height(title_row, title_row_height_twips)

    # Merge cells for title row.

    title_cell = title_row.cells[0]
    title_cell.merge(title_row.cells[1])

    # Set merged cell width to full table width.

    _set_cell_width(title_cell, content_width_twips)

    # Style title cell using config values.

    _set_cell_shading(title_cell, title_bg_hex)
    _set_cell_vertical_alignment(title_cell, "center")
    _add_title_text(title_cell, str(year), title_font_size, title_font_color)

    # === MONTH ROWS ===

    months = [
        ["January", "February"],
        ["March", "April"],
        ["May", "June"],
        ["July", "August"],
        ["September", "October"],
        ["November", "December"]
    ]

    for row_idx, month_row in enumerate(months):
        row = outer_table.rows[row_idx + 1]  # +1 to skip title row

        # Set explicit row height to fill available space.

        _set_row_height(row, month_row_height_twips)

        for col_idx, month_name in enumerate(month_row):
            cell = row.cells[col_idx]
            _set_cell_width(cell, cell_width_twips)

            month_num = row_idx * 2 + col_idx + 1
            _add_month_calendar(cell, year, month_num, month_name,
                               day_row_height_pt, month_name_gap_pt,
                               header_bg_hex)

    # Add borders around each cell.

    set_table_borders(outer_table, config.table)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_row_height
#
# Description:
#
#   Set explicit row height for a table row.
#
# Arguments:
#
#   row          : The table row.
#   height_twips : Height in twips (dxa).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_row_height(row, height_twips: int) -> None:

    # Set explicit row height for a table row.

    tr = row._tr
    tr_pr = tr.get_or_add_trPr()

    # Remove existing height if any.

    existing_height = tr_pr.find(qn('w:trHeight'))
    if existing_height is not None:
        tr_pr.remove(existing_height)

    # Set exact row height.

    tr_height = parse_xml(
        f'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{height_twips}" w:hRule="exact"/>'
    )
    tr_pr.append(tr_height)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_shading
#
# Description:
#
#   Set the background shading color of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   color_hex : Hex color string (e.g., "000000" for black).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str) -> None:

    # Set the background shading color of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tc_pr.append(shd)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_vertical_alignment
#
# Description:
#
#   Set the vertical alignment of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   alignment : Alignment value ("top", "center", "bottom").
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_vertical_alignment(cell, alignment: str) -> None:

    # Set the vertical alignment of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    # Remove existing vAlign if any.

    existing_valign = tc_pr.find(qn('w:vAlign'))
    if existing_valign is not None:
        tc_pr.remove(existing_valign)

    v_align = parse_xml(
        f'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{alignment}"/>'
    )
    tc_pr.append(v_align)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_title_text
#
# Description:
#
#   Add title text to a cell (bold, centered).
#
# Arguments:
#
#   cell       : The table cell.
#   text       : The text to add.
#   font_size  : Font size (Pt value).
#   font_color : Font color (RGBColor).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_title_text(cell, text: str, font_size, font_color) -> None:

    # Add title text to a cell (bold, centered).

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = font_size
    run.font.bold = True
    run.font.color.rgb = font_color


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_margins
#
# Description:
#
#   Set margins (padding) inside a table cell.
#
# Arguments:
#
#   cell   : The table cell.
#   top    : Top margin in twips.
#   bottom : Bottom margin in twips.
#   left   : Left margin in twips.
#   right  : Right margin in twips.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_margins(cell, top: int = 0, bottom: int = 0,
                      left: int = 0, right: int = 0) -> None:

    # Set margins (padding) inside a table cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = parse_xml(
        f'''<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>'''
    )
    tc_pr.append(tc_mar)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_month_calendar
#
# Description:
#
#   Add a mini calendar for a single month to a table cell.
#
# Arguments:
#
#   cell               : The table cell to add the calendar to.
#   year               : The year.
#   month              : The month number (1-12).
#   month_name         : The name of the month.
#   day_row_height_pt  : Height of day rows in points.
#   month_name_gap_pt  : Space between month name and days grid in points.
#   header_bg_hex      : Background color for month name (hex string, e.g., "BFBFBF").
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_month_calendar(cell, year: int, month: int, month_name: str,
                        day_row_height_pt: int, month_name_gap_pt: int,
                        header_bg_hex: str = None) -> None:

    # Add a mini calendar for a single month to a table cell.

    # Clear existing paragraph.

    cell.paragraphs[0].clear()

    # Add month name as header.

    header_para = cell.paragraphs[0]
    header_run = header_para.add_run(month_name)
    header_run.font.name = FONT_NAME
    header_run.font.size = Pt(10)
    header_run.font.bold = True
    header_run.font.color.rgb = COLOR_BLACK
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(month_name_gap_pt)

    # Add background shading to month name paragraph.

    if header_bg_hex:
        _set_paragraph_shading(header_para, header_bg_hex)

    # Get calendar data for the month.

    cal = cal_module.Calendar(firstweekday=6)  # Sunday first
    month_days = cal.monthdayscalendar(year, month)

    # Create mini table for days (header + weeks).

    num_weeks = len(month_days)
    mini_table = cell.add_table(rows=num_weeks + 1, cols=7)
    mini_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    mini_table.autofit = True

    # Row height for day rows (configurable vertical spacing).

    day_row_height = Pt(day_row_height_pt)

    # Day headers (S M T W T F S).

    day_headers = ["S", "M", "T", "W", "T", "F", "S"]
    header_row = mini_table.rows[0]
    header_row.height = day_row_height
    for i, day_header in enumerate(day_headers):
        header_cell = header_row.cells[i]
        para = header_cell.paragraphs[0]
        run = para.add_run(day_header)
        run.font.name = FONT_NAME
        run.font.size = Pt(7)
        run.font.bold = True
        run.font.color.rgb = COLOR_BLACK
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fill in the days.

    for week_idx, week in enumerate(month_days):
        week_row = mini_table.rows[week_idx + 1]
        week_row.height = day_row_height
        for day_idx, day in enumerate(week):
            day_cell = week_row.cells[day_idx]
            para = day_cell.paragraphs[0]
            if day != 0:
                run = para.add_run(str(day))
                run.font.name = FONT_NAME
                run.font.size = Pt(7)
                run.font.color.rgb = COLOR_BLACK
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove borders from mini table.

    _remove_table_borders(mini_table)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_paragraph_shading
#
# Description:
#
#   Set background shading for a paragraph.
#
# Arguments:
#
#   paragraph : The paragraph to shade.
#   color_hex : Hex color string (e.g., "BFBFBF" for gray).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_paragraph_shading(paragraph, color_hex: str) -> None:

    # Set background shading for a paragraph.

    p = paragraph._p
    p_pr = p.get_or_add_pPr()

    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    p_pr.append(shd)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _remove_table_borders
#
# Description:
#
#   Remove all borders from a table.
#
# Arguments:
#
#   table : The table to remove borders from.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _remove_table_borders(table) -> None:

    # Remove all borders from a table.

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    tbl_borders = parse_xml(
        '''<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="nil"/>
            <w:left w:val="nil"/>
            <w:bottom w:val="nil"/>
            <w:right w:val="nil"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
        </w:tblBorders>'''
    )

    tbl_pr.append(tbl_borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_layout_fixed
#
# Description:
#
#   Set the table layout to fixed to prevent auto-resizing.
#
# Arguments:
#
#   table : The table to set layout on.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_layout_fixed(table) -> None:

    # Set the table layout to fixed to prevent auto-resizing.

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tbl_pr)

    tbl_layout = parse_xml(
        '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="fixed"/>'
    )
    tbl_pr.append(tbl_layout)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_cell_margins
#
# Description:
#
#   Set default cell margins for the entire table.
#
# Arguments:
#
#   table  : The table to set margins on.
#   top    : Top margin in twips (None to keep default).
#   bottom : Bottom margin in twips (None to keep default).
#   left   : Left margin in twips (None to keep default).
#   right  : Right margin in twips (None to keep default).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_cell_margins(table, top: int = None, bottom: int = None,
                            left: int = None, right: int = None) -> None:

    # Set default cell margins for the entire table.

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tbl_pr)

    # Build tblCellMar element with only specified margins.

    margin_elements = []
    if top is not None:
        margin_elements.append(f'<w:top w:w="{top}" w:type="dxa"/>')
    if bottom is not None:
        margin_elements.append(f'<w:bottom w:w="{bottom}" w:type="dxa"/>')
    if left is not None:
        margin_elements.append(f'<w:left w:w="{left}" w:type="dxa"/>')
    if right is not None:
        margin_elements.append(f'<w:right w:w="{right}" w:type="dxa"/>')

    if margin_elements:
        tbl_cell_mar = parse_xml(
            f'<w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'{"".join(margin_elements)}'
            f'</w:tblCellMar>'
        )
        tbl_pr.append(tbl_cell_mar)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_grid
#
# Description:
#
#   Set the table grid column widths.
#
# Arguments:
#
#   table      : The table to set the grid on.
#   col_widths : List of column widths in twips (dxa).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_grid(table, col_widths: list[int]) -> None:

    # Set the table grid column widths.

    tbl = table._tbl

    # Remove existing grid if any.

    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    # Create new grid.

    grid_xml = '<w:tblGrid xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for width in col_widths:
        grid_xml += f'<w:gridCol w:w="{width}"/>'
    grid_xml += '</w:tblGrid>'

    tbl_grid = parse_xml(grid_xml)

    # Insert after tblPr.

    tbl_pr = tbl.tblPr
    if tbl_pr is not None:
        tbl_pr.addnext(tbl_grid)
    else:
        tbl.insert(0, tbl_grid)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_width
#
# Description:
#
#   Set the cell width in twips (dxa).
#
# Arguments:
#
#   cell         : The table cell.
#   width_twips  : Width in twips.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_width(cell, width_twips: int) -> None:

    # Set the cell width in twips (dxa).

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    # Remove existing width if any.

    existing_width = tc_pr.find(qn('w:tcW'))
    if existing_width is not None:
        tc_pr.remove(existing_width)

    tc_w = parse_xml(
        f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:w="{width_twips}" w:type="dxa"/>'
    )
    tc_pr.insert(0, tc_w)
