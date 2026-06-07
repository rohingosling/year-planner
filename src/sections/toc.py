#-----------------------------------------------------------------------------------------------------------------------
# Module:  toc.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Table of Contents generator for the Year Planner. Generates a custom table-based TOC serving as a hybrid TOC/Index,
#   pre-calculating page numbers for all content sections.
#-----------------------------------------------------------------------------------------------------------------------

import calendar
from dataclasses import dataclass
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from src.config import Config
from src.document import (
    add_page_break, add_config_info_overlay, get_content_width_twips,
    compute_table_row_height, MINIMIZED_PARAGRAPH_HEIGHT_TWIPS,
    get_title_row_height_twips, grayscale_to_hex, grayscale_to_rgb
)
from src.utils.styles import FONT_NAME, COLOR_BLACK


# Month names.

MONTH_NAMES = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December"
]

# Day names.

DAY_NAMES = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]


#-----------------------------------------------------------------------------------------------------------------------
# Class: TOCEntry
#
# Description:
#
#   A single entry in the Table of Contents.
#
# Attributes:
#
#   label         : Description text (e.g., "Goals", "January 1st, Thursday").
#   page_number   : Logical page number.
#   shading_level : 0 = none, 1 = section header, 2 = first item marker.
#-----------------------------------------------------------------------------------------------------------------------

@dataclass
class TOCEntry:
    label: str
    page_number: int
    shading_level: int = 0


#-----------------------------------------------------------------------------------------------------------------------
# Function: generate_toc
#
# Description:
#
#   Generate the Table of Contents section.
#
#   Creates a custom table-based TOC with entries for all numbered pages. The TOC fills pages and is double-sided.
#
# Arguments:
#
#   document : The Word document to add the TOC to.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def generate_toc(document: Document, config: Config) -> None:

    # Generate the Table of Contents section.

    # Get TOC config (with defaults).

    toc_config = config.raw.get('toc', {})
    rows_per_page = toc_config.get('rows_per_page', 40)

    # Build all TOC entries.

    entries = _build_toc_entries(config)

    # Calculate column widths once for all tables (consistency across pages).

    col_widths = _calculate_column_widths(entries, config)

    # Generate TOC pages.

    total_entries = len(entries)
    page_idx = 0

    for start_idx in range(0, total_entries, rows_per_page):

        # Add page break between TOC pages (not before first).

        if page_idx > 0:
            add_page_break(document, minimize_height=True)

        end_idx = min(start_idx + rows_per_page, total_entries)
        page_entries = entries[start_idx:end_idx]

        # Create TOC table for this page.
        # Last page: only create rows for actual entries, but keep same row height.

        is_last_page = (end_idx == total_entries)
        _create_toc_table(document, config, page_entries, rows_per_page, col_widths, is_last_page)

        # Determine if this is recto or verso (TOC starts on recto).

        is_recto = (page_idx % 2 == 0)
        add_config_info_overlay(document, config, is_recto=is_recto)

        page_idx += 1

    # Ensure TOC ends on verso so next section starts on recto.

    if page_idx % 2 == 1:

        # Odd number of pages means we ended on recto, add blank verso.

        add_page_break(document, minimize_height=True)

        # Add minimal paragraph to establish this as a page with body content.
        # (Overlay alone goes to header/footer, not body.)

        blank_para = document.add_paragraph()
        blank_para.paragraph_format.space_before = Pt(0)
        blank_para.paragraph_format.space_after = Pt(0)

        # Set minimal font size to minimize any visual impact.

        if blank_para.runs:
            blank_para.runs[0].font.size = Pt(1)
        add_config_info_overlay(document, config, is_recto=False, anchor_paragraph=blank_para)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _estimate_text_width_twips
#
# Description:
#
#   Estimate text width in twips based on character count and font size.
#
#   Uses average character width approximation for proportional fonts like Arial/Calibri.
#
# Arguments:
#
#   text         : The text string to measure.
#   font_size_pt : Font size in points.
#
# Returns:
#
#   Estimated width in twips.
#-----------------------------------------------------------------------------------------------------------------------

def _estimate_text_width_twips(text: str, font_size_pt: float) -> int:

    # Estimate text width in twips based on character count and font size.

    avg_char_width_pt = font_size_pt * 0.50  # Approximate for Arial/Calibri.
    text_width_pt = len(text) * avg_char_width_pt

    # Return data to caller.

    return int(text_width_pt * 20)  # 20 twips per point.


#-----------------------------------------------------------------------------------------------------------------------
# Function: _calculate_column_widths
#
# Description:
#
#   Calculate optimal column widths based on content.
#
#   - Label column: Width of longest label + padding.
#   - Page number column: Width of largest page number + padding.
#   - User entry column: Remaining space.
#
# Arguments:
#
#   entries : All TOC entries.
#   config  : Configuration with page settings.
#
# Returns:
#
#   Tuple of (label_width, entry_width, page_width) in twips.
#-----------------------------------------------------------------------------------------------------------------------

def _calculate_column_widths(entries: list[TOCEntry], config: Config) -> tuple[int, int, int]:

    # Calculate optimal column widths based on content.

    total_width = get_content_width_twips(config)
    font_size = 10  # Content font size in points.
    label_padding_twips = 150  # ~0.26cm padding for label column.
    page_padding_twips = 250  # ~0.44cm padding for page number column (needs more room).

    # Find longest label.

    max_label = max((e.label for e in entries), key=len, default="")
    label_width = _estimate_text_width_twips(max_label, font_size) + label_padding_twips

    # Find largest page number.

    max_page = max(e.page_number for e in entries)
    page_num_text = str(max_page)
    page_width = _estimate_text_width_twips(page_num_text, font_size) + page_padding_twips

    # User entry gets remaining space.

    entry_width = total_width - label_width - page_width

    # Return data to caller.

    return (label_width, entry_width, page_width)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _build_toc_entries
#
# Description:
#
#   Build all TOC entries with pre-calculated page numbers.
#
#   Page numbering starts at 1 on the Goals page.
#
#   Shading levels:
#   - Level 0: No shading (default).
#   - Level 1: Section headers (Goals, Backlog first, Week Planner first, month covers, Graph Paper first).
#   - Level 2: First item markers (first week of month in Week Planner, Mondays in daily spreads).
#
# Arguments:
#
#   config : Configuration with document settings.
#
# Returns:
#
#   List of TOCEntry objects for all numbered pages.
#-----------------------------------------------------------------------------------------------------------------------

def _build_toc_entries(config: Config) -> list[TOCEntry]:

    # Build all TOC entries with pre-calculated page numbers.

    entries = []
    current_page = 1
    year = config.document.year

    # === GOALS (pages 1-2: recto + blank verso) ===

    entries.append(TOCEntry("Goals", current_page, shading_level=1))
    current_page += 1
    entries.append(TOCEntry("", current_page))  # Blank verso.
    current_page += 1

    # === BACKLOG ===

    backlog_pages = config.raw.get('backlog', {}).get('page_count', 4)
    for i in range(backlog_pages):
        shading = 1 if i == 0 else 0  # Level 1 for first page only.
        entries.append(TOCEntry(f"Backlog ({i + 1}/{backlog_pages})", current_page, shading_level=shading))
        current_page += 1

    # === WEEK PLANNER ===

    weeks = _get_week_count(year)
    rows_per_page = config.raw.get('week_planner', {}).get('rows_per_page', 14)
    week_pages = (weeks + rows_per_page - 1) // rows_per_page  # Ceiling division.

    # Get weeks that contain the 1st of each month (for Level 2 shading).

    first_week_of_months = _get_first_weeks_of_months(year)

    for i in range(week_pages):
        start_week = i * rows_per_page + 1
        end_week = min((i + 1) * rows_per_page, weeks)

        # Check if any first-week-of-month falls within this page's weeks.

        has_first_week = any(start_week <= w <= end_week for w in first_week_of_months)

        if i == 0:
            shading = 1  # First page gets Level 1.
        elif has_first_week:
            shading = 2  # Pages with first-week-of-month get Level 2.
        else:
            shading = 0

        entries.append(TOCEntry(f"Week Planner (Weeks {start_week}-{end_week})", current_page, shading_level=shading))
        current_page += 1

    # === MONTHLY SECTIONS ===

    for month_num in range(1, 13):
        month_name = MONTH_NAMES[month_num - 1]
        num_days = calendar.monthrange(year, month_num)[1]

        # Month cover (recto) + blank verso - Level 1 shading for cover.

        entries.append(TOCEntry(f"{month_name}", current_page, shading_level=1))
        current_page += 1
        entries.append(TOCEntry("", current_page))  # Blank verso.
        current_page += 1

        # Daily spread pages (2 tables per page side).

        days = [date(year, month_num, day) for day in range(1, num_days + 1)]
        day_idx = 0

        while day_idx < len(days):

            # First day on this page.

            day1 = days[day_idx]
            label1 = _format_day_label(day1)

            # Level 2 shading for Saturday (5) and Sunday (6).

            shading1 = 2 if day1.weekday() in (5, 6) else 0
            entries.append(TOCEntry(label1, current_page, shading_level=shading1))
            day_idx += 1

            # Second day on this page (if exists).

            if day_idx < len(days):
                day2 = days[day_idx]
                label2 = _format_day_label(day2)

                # Level 2 shading for Saturday (5) and Sunday (6).

                shading2 = 2 if day2.weekday() in (5, 6) else 0
                entries.append(TOCEntry(label2, current_page, shading_level=shading2))
                day_idx += 1

            current_page += 1

        # Daily spread always ends on verso, check if we need blank verso.

        num_page_sides = (num_days + 1) // 2
        if num_page_sides % 2 == 1:

            # Ended on recto, blank verso was added.

            entries.append(TOCEntry("", current_page))
            current_page += 1

    # === TERMS AND DEFINITIONS ===

    td_pages = config.raw.get('terms_definitions', {}).get('page_count', 4)
    for i in range(td_pages):
        shading = 1 if i == 0 else 0  # Level 1 for first page only.
        entries.append(TOCEntry(f"Terms and Definitions ({i + 1}/{td_pages})", current_page, shading_level=shading))
        current_page += 1

    # === GRAPH PAPER ===

    graph_pages = config.raw.get('graph_paper', {}).get('page_count', 12)
    for i in range(graph_pages):
        shading = 1 if i == 0 else 0  # Level 1 for first page only.
        entries.append(TOCEntry(f"Graph Paper ({i + 1}/{graph_pages})", current_page, shading_level=shading))
        current_page += 1
        entries.append(TOCEntry("", current_page))  # Blank verso.
        current_page += 1

    # Return data to caller.

    return entries


#-----------------------------------------------------------------------------------------------------------------------
# Function: _get_first_weeks_of_months
#
# Description:
#
#   Get the ISO week numbers that contain the 1st day of each month.
#
# Arguments:
#
#   year : The year.
#
# Returns:
#
#   Set of ISO week numbers (1-53) that contain a month's first day.
#-----------------------------------------------------------------------------------------------------------------------

def _get_first_weeks_of_months(year: int) -> set[int]:

    # Get the ISO week numbers that contain the 1st day of each month.

    first_weeks = set()
    for month in range(1, 13):
        first_day = date(year, month, 1)
        iso_week = first_day.isocalendar()[1]
        first_weeks.add(iso_week)

    # Return data to caller.

    return first_weeks


#-----------------------------------------------------------------------------------------------------------------------
# Function: _get_week_count
#
# Description:
#
#   Get the number of ISO 8601 weeks in a year (52 or 53).
#
# Arguments:
#
#   year : The year.
#
# Returns:
#
#   Number of weeks (52 or 53).
#-----------------------------------------------------------------------------------------------------------------------

def _get_week_count(year: int) -> int:

    # Get the number of ISO 8601 weeks in a year (52 or 53).

    # A year has 53 weeks if:
    # - January 1st is Thursday, OR
    # - December 31st is Thursday.
    # (In leap years, either of these conditions means 53 weeks.)

    jan1 = date(year, 1, 1)
    dec31 = date(year, 12, 31)

    # weekday() returns 0=Monday, 3=Thursday.

    if jan1.weekday() == 3 or dec31.weekday() == 3:
        return 53
    return 52


#-----------------------------------------------------------------------------------------------------------------------
# Function: _format_day_label
#
# Description:
#
#   Format a day for the TOC entry.
#
#   Format: "Week N, Month Nth, Day" (e.g., "Week 1, January 1st, Thursday").
#
# Arguments:
#
#   day : The date.
#
# Returns:
#
#   Formatted label string.
#-----------------------------------------------------------------------------------------------------------------------

def _format_day_label(day: date) -> str:

    # Format a day for the TOC entry.

    week_num = day.isocalendar()[1]
    month_name = MONTH_NAMES[day.month - 1]
    day_name = DAY_NAMES[day.weekday()]
    ordinal = _get_ordinal_suffix(day.day)

    # Return data to caller.

    return f"Week {week_num}, {month_name} {day.day}{ordinal}, {day_name}"


#-----------------------------------------------------------------------------------------------------------------------
# Function: _get_ordinal_suffix
#
# Description:
#
#   Get ordinal suffix for a number (st, nd, rd, th).
#
# Arguments:
#
#   n : The number.
#
# Returns:
#
#   Ordinal suffix string.
#-----------------------------------------------------------------------------------------------------------------------

def _get_ordinal_suffix(n: int) -> str:

    # Get ordinal suffix for a number (st, nd, rd, th).

    if 11 <= n <= 13:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")


#-----------------------------------------------------------------------------------------------------------------------
# Function: _create_toc_table
#
# Description:
#
#   Create a TOC table for one page.
#
#   Structure:
#   - Title row: "Table of Contents" spanning all columns.
#   - Content rows: Page Label | User Entry | Page Number.
#
# Arguments:
#
#   document      : The Word document.
#   config        : Configuration settings.
#   entries       : TOC entries for this page.
#   rows_per_page : Number of content rows per page (used for row height calculation).
#   col_widths    : Tuple of (label_width, entry_width, page_width) in twips.
#   is_last_page  : If True, only create rows for actual entries (no empty rows).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _create_toc_table(document: Document, config: Config,
                      entries: list[TOCEntry], rows_per_page: int,
                      col_widths: tuple[int, int, int],
                      is_last_page: bool = False) -> None:

    # Create a TOC table for one page.

    label_width, entry_width, page_width = col_widths
    total_width_twips = get_content_width_twips(config)

    # Create table: title row + content rows.
    # On last page, only create rows for actual entries (no empty rows).

    actual_content_rows = len(entries) if is_last_page else rows_per_page
    total_rows = 1 + actual_content_rows
    table = document.add_table(rows=total_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set fixed table layout.

    _set_table_layout_fixed(table)
    _set_table_grid(table, list(col_widths))

    # Get row heights.

    title_row_height_twips = get_title_row_height_twips(config)

    # Compute content row height dynamically.

    content_row_height_twips = compute_table_row_height(
        config=config,
        num_content_rows=rows_per_page,
        title_row_height_twips=title_row_height_twips,
        header_row_height_twips=0,  # No header row.
        preceding_paragraph_height_twips=MINIMIZED_PARAGRAPH_HEIGHT_TWIPS
    )

    # Get title row styling from config.

    title_bg_hex = grayscale_to_hex(config.table.title_row.background_grayscale)
    title_font_rgb = grayscale_to_rgb(config.table.title_row.font_grayscale)
    title_font_color = RGBColor(*title_font_rgb)
    title_font_size = Pt(config.table.title_row.font_size)

    # Get TOC shading config.

    toc_config = config.raw.get('toc', {})
    section_grayscale = toc_config.get('section_grayscale', 15)
    first_item_grayscale = toc_config.get('first_item_grayscale', 5)

    # === TITLE ROW ===

    title_row = table.rows[0]
    _set_row_height(title_row, title_row_height_twips)

    # Merge all cells in title row.

    title_cell = title_row.cells[0]
    for i in range(1, 3):
        title_cell.merge(title_row.cells[i])

    _set_cell_width(title_cell, total_width_twips)
    _set_cell_shading(title_cell, title_bg_hex)
    _set_cell_vertical_alignment(title_cell, "center")
    _add_cell_text(title_cell, "Table of Contents", size=title_font_size,
                   bold=True, color=title_font_color, align=None)

    # === CONTENT ROWS ===
    # On last page, we only have actual_content_rows (no empty rows).
    # Row height is still based on rows_per_page for consistency.

    for row_idx in range(actual_content_rows):
        row = table.rows[row_idx + 1]
        _set_row_height(row, content_row_height_twips)

        # Get entry for this row.

        entry = entries[row_idx]
        label_text = entry.label
        page_text = str(entry.page_number)  # Always show page number.
        shading_level = entry.shading_level

        # Label cell (left-aligned).

        label_cell = row.cells[0]
        _set_cell_width(label_cell, label_width)
        _set_cell_vertical_alignment(label_cell, "center")
        if label_text:
            _add_cell_text(label_cell, label_text, size=Pt(10), bold=False,
                           color=COLOR_BLACK, align=WD_ALIGN_PARAGRAPH.LEFT)

        # User entry cell (blank).

        entry_cell = row.cells[1]
        _set_cell_width(entry_cell, entry_width)
        _set_cell_vertical_alignment(entry_cell, "center")

        # Page number cell (right-aligned).

        page_cell = row.cells[2]
        _set_cell_width(page_cell, page_width)
        _set_cell_vertical_alignment(page_cell, "center")
        if page_text:
            _add_cell_text(page_cell, page_text, size=Pt(10), bold=False,
                           color=COLOR_BLACK, align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Apply shading based on level.

        if shading_level == 1:
            shading_hex = grayscale_to_hex(section_grayscale)
            _set_cell_shading(label_cell, shading_hex)
            _set_cell_shading(entry_cell, shading_hex)
            _set_cell_shading(page_cell, shading_hex)
        elif shading_level == 2:
            shading_hex = grayscale_to_hex(first_item_grayscale)
            _set_cell_shading(label_cell, shading_hex)
            _set_cell_shading(entry_cell, shading_hex)
            _set_cell_shading(page_cell, shading_hex)

    # Apply borders.

    _set_table_borders(table, config)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_layout_fixed
#
# Description:
#
#   Set table layout to fixed.
#
# Arguments:
#
#   table : The table to set layout on.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_layout_fixed(table) -> None:

    # Set table layout to fixed.

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tbl_pr)

    tbl_layout = parse_xml(
        '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="fixed"/>'
    )
    tbl_pr.append(tbl_layout)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_grid
#
# Description:
#
#   Set the table grid column widths.
#
# Arguments:
#
#   table      : The table to set the grid on.
#   col_widths : List of column widths in dxa (twips).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_grid(table, col_widths: list[int]) -> None:

    # Set the table grid column widths.

    tbl = table._tbl

    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    grid_xml = '<w:tblGrid xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for width in col_widths:
        grid_xml += f'<w:gridCol w:w="{width}"/>'
    grid_xml += '</w:tblGrid>'

    tbl_grid = parse_xml(grid_xml)

    tbl_pr = tbl.tblPr
    if tbl_pr is not None:
        tbl_pr.addnext(tbl_grid)
    else:
        tbl.insert(0, tbl_grid)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_row_height
#
# Description:
#
#   Set the row height in twips.
#
# Arguments:
#
#   row          : The table row.
#   height_twips : Height in twips (dxa).
#   exact        : If True, row height is exactly the specified value (won't grow).
#                  If False, row height is at least the specified value (can grow).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_row_height(row, height_twips: int, exact: bool = True) -> None:

    # Set the row height in twips.

    tr = row._tr
    tr_pr = tr.get_or_add_trPr()

    existing_height = tr_pr.find(qn('w:trHeight'))
    if existing_height is not None:
        tr_pr.remove(existing_height)

    h_rule = "exact" if exact else "atLeast"
    tr_height = parse_xml(
        f'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{height_twips}" w:hRule="{h_rule}"/>'
    )
    tr_pr.append(tr_height)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_width
#
# Description:
#
#   Set cell width in dxa (twips).
#
# Arguments:
#
#   cell      : The table cell.
#   width_dxa : Width in dxa.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_width(cell, width_dxa: int) -> None:

    # Set cell width in dxa (twips).

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    existing_width = tc_pr.find(qn('w:tcW'))
    if existing_width is not None:
        tc_pr.remove(existing_width)

    tc_w = parse_xml(
        f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:w="{width_dxa}" w:type="dxa"/>'
    )
    tc_pr.insert(0, tc_w)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_vertical_alignment
#
# Description:
#
#   Set vertical alignment of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   alignment : Alignment value ("top", "center", "bottom").
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_vertical_alignment(cell, alignment: str) -> None:

    # Set vertical alignment of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    existing_valign = tc_pr.find(qn('w:vAlign'))
    if existing_valign is not None:
        tc_pr.remove(existing_valign)

    v_align = parse_xml(
        f'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{alignment}"/>'
    )
    tc_pr.append(v_align)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_shading
#
# Description:
#
#   Set the background shading color of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   color_hex : Hex color string (e.g., "000000" for black).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str) -> None:

    # Set the background shading color of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tc_pr.append(shd)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_cell_text
#
# Description:
#
#   Add formatted text to a table cell.
#
# Arguments:
#
#   cell  : The table cell.
#   text  : The text to add.
#   size  : Font size (None to inherit from the document default).
#   bold  : Whether to make the text bold.
#   color : Font color (RGBColor).
#   align : Paragraph alignment (None for default LEFT).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_cell_text(cell, text: str, size=None, bold: bool = False,
                   color=COLOR_BLACK, align=None) -> None:

    # Add formatted text to a table cell.

    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = FONT_NAME
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = bold
    run.font.color.rgb = color


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_borders
#
# Description:
#
#   Set table borders using config settings.
#
# Arguments:
#
#   table  : The table to set borders on.
#   config : Configuration with border settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_borders(table, config: Config) -> None:

    # Set table borders using config settings.

    border_color_hex = grayscale_to_hex(config.table.border.grayscale)
    border_size = int(config.table.border.thickness * 8)

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    tbl_borders = parse_xml(
        f'''<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:left w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:right w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideH w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideV w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
        </w:tblBorders>'''
    )

    tbl_pr.append(tbl_borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)
