#-----------------------------------------------------------------------------------------------------------------------
# Module:  goals.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Goals page generator for the Year Planner. Generates a single page with a Goals table for long-term goal planning.
#-----------------------------------------------------------------------------------------------------------------------

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from src.config import Config
from src.document import (
    add_page_break, get_content_width_twips, compute_table_row_height,
    MINIMIZED_PARAGRAPH_HEIGHT_TWIPS, TWIPS_PER_CM, add_config_info_overlay,
    get_title_row_height_twips, grayscale_to_hex, grayscale_to_rgb
)
from src.utils.styles import FONT_NAME, COLOR_BLACK


#-----------------------------------------------------------------------------------------------------------------------
# Function: generate_goals_page
#
# Description:
#
#   Generate the Goals page.
#
#   Creates a single page with a Goals table. The table fills the page vertically with configurable columns and rows.
#   Starts on the recto, followed by a blank verso.
#
# Arguments:
#
#   document : The Word document to add the goals page to.
#   config   : Configuration with document settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def generate_goals_page(document: Document, config: Config) -> None:

    # Generate the Goals page.

    # Get goals config.

    goals_config = config.raw.get('goals', {})
    num_columns = goals_config.get('columns', 2)
    num_rows = goals_config.get('rows', 4)

    # Create the goals table.

    _create_goals_table(document, config, num_columns, num_rows)

    # Add overlay (recto page).

    add_config_info_overlay(document, config, is_recto=True)

    # Add blank verso page so the next section starts on the recto.
    # Use minimize_height=True to avoid overflow since the table fills the page.

    add_page_break(document, minimize_height=True)
    add_config_info_overlay(document, config, is_recto=False)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _create_goals_table
#
# Description:
#
#   Create the Goals table.
#
#   Structure:
#   - Title row: "Goals" spanning all columns.
#   - Content rows: Empty cells for the user to fill in.
#
# Arguments:
#
#   document    : The Word document.
#   config      : Configuration settings.
#   num_columns : Number of columns in the table.
#   num_rows    : Number of content rows.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _create_goals_table(document: Document, config: Config,
                        num_columns: int, num_rows: int) -> None:

    # Create the Goals table.

    # Create table: title row + content rows.

    total_rows = 1 + num_rows  # title + content
    table = document.add_table(rows=total_rows, cols=num_columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set fixed table layout.

    _set_table_layout_fixed(table)

    # Calculate column widths (equal distribution).

    total_width_twips = get_content_width_twips(config)
    col_width_twips = total_width_twips // num_columns
    col_widths = [col_width_twips] * num_columns

    # Adjust last column for any rounding.

    col_widths[-1] = total_width_twips - (col_width_twips * (num_columns - 1))

    # Set table grid column widths.

    _set_table_grid(table, col_widths)

    # Get row heights from config.

    title_row_height_twips = get_title_row_height_twips(config)

    # Compute content row height dynamically to fill the exact vertical content area.
    # Goals table has no header row, only a title row.

    content_row_height_twips = compute_table_row_height(
        config=config,
        num_content_rows=num_rows,
        title_row_height_twips=title_row_height_twips,
        header_row_height_twips=0,  # No header row
        preceding_paragraph_height_twips=MINIMIZED_PARAGRAPH_HEIGHT_TWIPS
    )

    # Get title row styling from config.

    title_bg_hex = grayscale_to_hex(config.table.title_row.background_grayscale)
    title_font_rgb = grayscale_to_rgb(config.table.title_row.font_grayscale)
    title_font_color = RGBColor(*title_font_rgb)
    title_font_size = Pt(config.table.title_row.font_size)

    # === TITLE ROW ===

    title_row = table.rows[0]
    _set_row_height(title_row, title_row_height_twips)

    # Merge all cells in the title row.

    title_cell = title_row.cells[0]
    for i in range(1, num_columns):
        title_cell.merge(title_row.cells[i])

    # Set title cell width.

    _set_cell_width(title_cell, total_width_twips)

    # Title cell: "Goals" - config background, config font, bold.

    _set_cell_shading(title_cell, title_bg_hex)
    _set_cell_vertical_alignment(title_cell, "center")
    _add_cell_text(title_cell, "Goals", size=title_font_size, bold=True,
                   color=title_font_color, align=None)

    # === CONTENT ROWS ===

    for row_idx in range(num_rows):
        row = table.rows[row_idx + 1]
        _set_row_height(row, content_row_height_twips)

        # Set cell widths and alignment for each column.

        for col_idx in range(num_columns):
            cell = row.cells[col_idx]
            _set_cell_width(cell, col_widths[col_idx])
            _set_cell_vertical_alignment(cell, "center")
            # Leave cells empty for the user to fill in.

    # Apply borders using config settings.

    _set_table_borders(table, config)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_layout_fixed
#
# Description:
#
#   Set the table layout to fixed so columns do not auto-resize.
#
# Arguments:
#
#   table : The table to set layout on.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_layout_fixed(table) -> None:

    # Set the table layout to fixed.

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(
            r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tbl_pr)

    tbl_layout = parse_xml(
        '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="fixed"/>'
    )
    tbl_pr.append(tbl_layout)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_grid
#
# Description:
#
#   Set the table grid column widths.
#
# Arguments:
#
#   table      : The table to set the grid on.
#   col_widths : List of column widths in dxa (twips).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_grid(table, col_widths: list[int]) -> None:

    # Set the table grid column widths.

    tbl = table._tbl

    # Remove the existing grid if any.

    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)

    # Create a new grid.

    grid_xml = '<w:tblGrid xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    for width in col_widths:
        grid_xml += f'<w:gridCol w:w="{width}"/>'
    grid_xml += '</w:tblGrid>'

    tbl_grid = parse_xml(grid_xml)

    # Insert after tblPr.

    tbl_pr = tbl.tblPr
    if tbl_pr is not None:
        tbl_pr.addnext(tbl_grid)
    else:
        tbl.insert(0, tbl_grid)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_row_height
#
# Description:
#
#   Set the row height in twips.
#
# Arguments:
#
#   row          : The table row.
#   height_twips : Height in twips (dxa).
#   exact        : If True, row height is exactly the specified value (won't grow).
#                  If False, row height is at least the specified value (can grow).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_row_height(row, height_twips: int, exact: bool = True) -> None:

    # Set the row height in twips.

    tr = row._tr
    tr_pr = tr.get_or_add_trPr()

    # Remove the existing height if any.

    existing_height = tr_pr.find(qn('w:trHeight'))
    if existing_height is not None:
        tr_pr.remove(existing_height)

    h_rule = "exact" if exact else "atLeast"
    tr_height = parse_xml(
        f'<w:trHeight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{height_twips}" w:hRule="{h_rule}"/>'
    )
    tr_pr.append(tr_height)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_width
#
# Description:
#
#   Set the cell width in dxa (twips).
#
# Arguments:
#
#   cell      : The table cell.
#   width_dxa : Width in dxa.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_width(cell, width_dxa: int) -> None:

    # Set the cell width in dxa (twips).

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    # Remove the existing width if any.

    existing_width = tc_pr.find(qn('w:tcW'))
    if existing_width is not None:
        tc_pr.remove(existing_width)

    tc_w = parse_xml(
        f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:w="{width_dxa}" w:type="dxa"/>'
    )
    tc_pr.insert(0, tc_w)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_vertical_alignment
#
# Description:
#
#   Set the vertical alignment of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   alignment : Alignment value ("top", "center", "bottom").
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_vertical_alignment(cell, alignment: str) -> None:

    # Set the vertical alignment of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    # Remove the existing vAlign if any.

    existing_valign = tc_pr.find(qn('w:vAlign'))
    if existing_valign is not None:
        tc_pr.remove(existing_valign)

    v_align = parse_xml(
        f'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="{alignment}"/>'
    )
    tc_pr.append(v_align)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_cell_shading
#
# Description:
#
#   Set the background shading color of a cell.
#
# Arguments:
#
#   cell      : The table cell.
#   color_hex : Hex color string (e.g., "000000" for black).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str) -> None:

    # Set the background shading color of a cell.

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tc_pr.append(shd)


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_cell_text
#
# Description:
#
#   Add formatted text to a table cell.
#
# Arguments:
#
#   cell  : The table cell.
#   text  : The text to add.
#   size  : Font size (None to inherit from the document default).
#   bold  : Whether to make the text bold.
#   color : Font color (RGBColor).
#   align : Paragraph alignment (None for default LEFT).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_cell_text(cell, text: str, size=None, bold: bool = False,
                   color=COLOR_BLACK, align=None) -> None:

    # Add formatted text to a table cell.

    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = FONT_NAME
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = bold
    run.font.color.rgb = color


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_table_borders
#
# Description:
#
#   Set table borders using config settings.
#
# Arguments:
#
#   table  : The table to set borders on.
#   config : Configuration with border settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_table_borders(table, config: Config) -> None:

    # Set table borders using config settings.

    # Get border settings from config.

    border_color_hex = grayscale_to_hex(config.table.border.grayscale)
    border_size = int(config.table.border.thickness * 8)  # Eighths of a point

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    tbl_borders = parse_xml(
        f'''<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:left w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:right w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideH w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
            <w:insideV w:val="single" w:sz="{border_size}" w:color="{border_color_hex}"/>
        </w:tblBorders>'''
    )

    tbl_pr.append(tbl_borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)
