#-----------------------------------------------------------------------------------------------------------------------
# Module:  styles.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Style definitions for Year Planner document. Provides consistent font and paragraph styling throughout
#   the document.
#-----------------------------------------------------------------------------------------------------------------------

from docx                import Document
from docx.shared         import Pt, RGBColor
from docx.enum.text      import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.table          import _Cell

# Global constants.

FONT_NAME          = "Times New Roman"
FONT_SIZE_TITLE    = Pt ( 36 )
FONT_SIZE_SUBTITLE = Pt ( 14 )
FONT_SIZE_NORMAL   = Pt ( 11 )
FONT_SIZE_SMALL    = Pt ( 9 )

COLOR_BLACK = RGBColor ( 0, 0, 0 )


#-----------------------------------------------------------------------------------------------------------------------
# Function: apply_title_style
#
# Description:
#
#   Apply title styling to a paragraph.
#
# Arguments:
#
#   paragraph : The paragraph to style.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def apply_title_style ( paragraph: Paragraph ) -> None:

    # Apply title styling to a paragraph.

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run                 = paragraph.runs [ 0 ] if paragraph.runs else paragraph.add_run ()
    run.font.name       = FONT_NAME
    run.font.size       = FONT_SIZE_TITLE
    run.font.bold       = True
    run.font.color.rgb  = COLOR_BLACK


#-----------------------------------------------------------------------------------------------------------------------
# Function: apply_subtitle_style
#
# Description:
#
#   Apply subtitle styling to a paragraph.
#
# Arguments:
#
#   paragraph : The paragraph to style.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def apply_subtitle_style ( paragraph: Paragraph ) -> None:

    # Apply subtitle styling to a paragraph.

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run                 = paragraph.runs [ 0 ] if paragraph.runs else paragraph.add_run ()
    run.font.name       = FONT_NAME
    run.font.size       = FONT_SIZE_SUBTITLE
    run.font.bold       = False
    run.font.color.rgb  = COLOR_BLACK


#-----------------------------------------------------------------------------------------------------------------------
# Function: apply_normal_style
#
# Description:
#
#   Apply normal text styling to a paragraph.
#
# Arguments:
#
#   paragraph : The paragraph to style.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def apply_normal_style ( paragraph: Paragraph ) -> None:

    # Apply normal text styling to a paragraph.

    run                = paragraph.runs [ 0 ] if paragraph.runs else paragraph.add_run ()
    run.font.name      = FONT_NAME
    run.font.size      = FONT_SIZE_NORMAL
    run.font.color.rgb = COLOR_BLACK


#-----------------------------------------------------------------------------------------------------------------------
# Function: apply_cell_style
#
# Description:
#
#   Apply styling to a table cell.
#
# Arguments:
#
#   cell      : The table cell to style.
#   bold      : Whether to make the text bold.
#   alignment : Text alignment within the cell.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def apply_cell_style ( cell: _Cell, bold: bool = False,
                       alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT ) -> None:

    # Apply styling to a table cell.

    paragraph           = cell.paragraphs [ 0 ]
    paragraph.alignment = alignment
    run                 = paragraph.runs [ 0 ] if paragraph.runs else paragraph.add_run ()
    run.font.name       = FONT_NAME
    run.font.size       = FONT_SIZE_NORMAL
    run.font.bold       = bold
    run.font.color.rgb  = COLOR_BLACK
