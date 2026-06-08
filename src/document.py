#-----------------------------------------------------------------------------------------------------------------------
# Module:  document.py
# Project: Year Planner Generator
# Version: 1.1
# Author:  Rohin Gosling
#
# Description:
#
#   Document initialization and configuration for the Year Planner. Handles document setup, page layout, and margins.
#-----------------------------------------------------------------------------------------------------------------------

from docx              import Document
from docx.shared       import Cm, Pt, Twips
from docx.enum.section import WD_ORIENT
from docx.enum.style   import WD_STYLE_TYPE
from docx.enum.text    import WD_LINE_SPACING
from docx.section      import Section
from docx.oxml         import parse_xml
from docx.oxml.ns      import qn

from src.config import Config

# Global constants.

# Conversion constant: twips per centimeter.
# 1 inch = 1440 twips, 1 inch = 2.54 cm, so 1 cm = 1440/2.54 ≈ 566.93 twips.

TWIPS_PER_CM = 1440 / 2.54

# Conversion constant: twips per point.
# 1 point = 1/72 inch, 1 inch = 1440 twips, so 1 point = 20 twips.

TWIPS_PER_PT = 20

# Default document settings.

DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_FONT_SIZE = Pt ( 11 )


#-----------------------------------------------------------------------------------------------------------------------
# Function: create_document
#
# Description:
#
#   Create and initialize a new Word document with configured page settings.
#
# Arguments:
#
#   config : Configuration object with page settings.
#
# Returns:
#
#   Initialized Document object.
#-----------------------------------------------------------------------------------------------------------------------

def create_document ( config: Config ) -> Document:

    # Create and initialize a new Word document with configured page settings.

    document = Document ()

    # Configure default styles.

    _configure_default_styles ( document )

    # Enable mirror margins for duplex printing.
    # This makes Word use "Inside/Outside" margins instead of "Left/Right" and automatically alternates the gutter
    # position for recto/verso pages.

    _enable_mirror_margins ( document )

    # Configure the default section.

    section = document.sections [ 0 ]
    configure_section ( section, config )

    # Return data to caller.

    return document


#-----------------------------------------------------------------------------------------------------------------------
# Function: _enable_mirror_margins
#
# Description:
#
#   Enable mirror margins for duplex printing.
#
#   When enabled, Word uses "Inside" and "Outside" margins instead of "Left" and "Right", and the gutter automatically
#   alternates between recto (right-hand, odd) and verso (left-hand, even) pages.
#
# Arguments:
#
#   document : The document to configure.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _enable_mirror_margins ( document: Document ) -> None:

    # Enable mirror margins for duplex printing.

    # Access the document settings.

    settings = document.settings.element

    # Add the mirrorMargins element.

    mirror_margins = parse_xml (
        '<w:mirrorMargins xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    settings.append ( mirror_margins )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _configure_default_styles
#
# Description:
#
#   Configure document-wide default styles.
#
#   Sets:
#   - Default font: Times New Roman, 11pt.
#   - Default paragraph spacing: Before=0, After=0.
#   - Default line spacing: Single.
#
# Arguments:
#
#   document : The document to configure.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _configure_default_styles ( document: Document ) -> None:

    # Configure document-wide default styles.

    # Configure the Normal style (affects all paragraphs by default).

    style                                    = document.styles [ 'Normal' ]
    style.font.name                          = DEFAULT_FONT_NAME
    style.font.size                          = DEFAULT_FONT_SIZE
    style.paragraph_format.space_before      = Pt ( 0 )
    style.paragraph_format.space_after       = Pt ( 0 )
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


#-----------------------------------------------------------------------------------------------------------------------
# Function: configure_section
#
# Description:
#
#   Configure page layout for a document section.
#
#   With mirror margins enabled, Word automatically handles gutter positioning for duplex printing:
#   - Recto (odd) pages: gutter on left (inside/binding edge).
#   - Verso (even) pages: gutter on right (inside/binding edge).
#
# Arguments:
#
#   section : The section to configure.
#   config  : Configuration with page settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def configure_section ( section: Section, config: Config ) -> None:

    # Configure page layout for a document section.

    # Set page size (A4: 21.0 x 29.7 cm).

    section.page_width  = Cm ( config.page.width )
    section.page_height = Cm ( config.page.height )
    section.orientation = WD_ORIENT.PORTRAIT

    # Set margins (in cm).
    # With mirror margins: left = inside (binding), right = outside.

    section.top_margin    = Cm ( config.page.margin_top )
    section.bottom_margin = Cm ( config.page.margin_bottom )
    section.left_margin   = Cm ( config.page.margin_left )    # Inside margin (binding edge)
    section.right_margin  = Cm ( config.page.margin_right )  # Outside margin (outer edge)

    # Gutter adds extra space on the binding edge for duplex printing.
    # With mirror margins, Word automatically alternates the gutter side.

    section.gutter = Cm ( config.page.gutter_size )

    # Footer distance from bottom.

    section.footer_distance = Cm ( config.page.page_number_position )

    # Header distance from top (set to 0 so headers don't affect the content area).

    section.header_distance = Cm ( 0 )

    # Add debug visualization if enabled.

    if config.debug.enabled:
        _add_debug_visualization ( section, config )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_debug_visualization
#
# Description:
#
#   Add debug visualization using odd/even headers for correct recto/verso positioning.
#
#   Draws:
#   - Blue vertical line at gutter boundary (gutter distance from binding edge).
#   - Red rectangle at content area boundary (gutter + margin from binding edge).
#   - Green horizontal lines at header/footer text region boundaries.
#
#   Layout:
#   - Recto: │←gutter→●←margin→■ content ■←margin→│
#   - Verso: │←margin→■ content ■←margin→●←gutter→│
#
# Arguments:
#
#   section : The section to add visualization to.
#   config  : Configuration with page settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_debug_visualization ( section: Section, config: Config ) -> None:

    # Add debug visualization using odd/even headers for correct recto/verso positioning.

    # Enable different odd/even headers.

    sect_pr      = section._sectPr
    even_and_odd = sect_pr.find ( qn ( 'w:evenAndOddHeaders' ) )

    if even_and_odd is None:

        even_and_odd = parse_xml (
            '<w:evenAndOddHeaders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        sect_pr.append ( even_and_odd )

    # Get measurements.

    margin_top_cm      = config.page.margin_top
    margin_bottom_cm   = config.page.margin_bottom
    margin_left_cm     = config.page.margin_left
    margin_right_cm    = config.page.margin_right
    gutter_cm          = config.page.gutter_size
    page_width_cm      = config.page.width
    page_height_cm     = config.page.height
    footer_distance_cm = config.page.page_number_position

    # Content area boundary (where the red rectangle goes) = gutter + margin from binding edge.

    content_left_cm   = gutter_cm + margin_left_cm  # Binding side (gutter + left margin)
    content_right_cm  = margin_right_cm  # Non-binding side (right margin only)
    content_top_cm    = margin_top_cm
    content_bottom_cm = margin_bottom_cm

    # Green line positions (header/footer text boundaries).
    # Top green line: at top margin (boundary between header and content).
    # Bottom green line: footer distance from bottom (where footer text region starts).

    green_top_y_cm    = margin_top_cm
    green_bottom_y_cm = page_height_cm - footer_distance_cm

    # Recto (odd) pages: binding on left.

    recto_red_left  = content_left_cm
    recto_red_right = page_width_cm - content_right_cm
    recto_blue_x    = gutter_cm

    # Verso (even) pages: binding on right.

    verso_red_left  = content_right_cm
    verso_red_right = page_width_cm - content_left_cm
    verso_blue_x    = page_width_cm - gutter_cm

    # Create the odd page header (recto).

    odd_header                       = section.header
    odd_header.is_linked_to_previous = False

    _add_debug_shapes_to_header (
        odd_header,
        recto_red_left, content_top_cm,
        recto_red_right, page_height_cm - content_bottom_cm,
        recto_blue_x, page_height_cm, page_width_cm,
        green_top_y_cm, green_bottom_y_cm
    )

    # Create the even page header (verso).

    even_header                       = section.even_page_header
    even_header.is_linked_to_previous = False

    _add_debug_shapes_to_header (
        even_header,
        verso_red_left, content_top_cm,
        verso_red_right, page_height_cm - content_bottom_cm,
        verso_blue_x, page_height_cm, page_width_cm,
        green_top_y_cm, green_bottom_y_cm
    )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_debug_shapes_to_header
#
# Description:
#
#   Add debug visualization shapes to a header.
#
#   Draws:
#   - Red rectangle for content area boundary.
#   - Blue vertical line for gutter boundary.
#   - Green horizontal lines for header/footer text region boundaries.
#
# Arguments:
#
#   header           : The header to add shapes to.
#   red_left_cm      : Left edge of red rectangle from page left (cm).
#   red_top_cm       : Top edge of red rectangle from page top (cm).
#   red_right_cm     : Right edge of red rectangle from page left (cm).
#   red_bottom_cm    : Bottom edge of red rectangle from page top (cm).
#   blue_x_cm        : X position of blue line from page left (cm).
#   page_height_cm   : Full page height for blue line (cm).
#   page_width_cm    : Full page width for green lines (cm).
#   green_top_y_cm   : Y position of top green line (header boundary) from page top (cm).
#   green_bottom_y_cm : Y position of bottom green line (footer boundary) from page top (cm).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_debug_shapes_to_header (
    header,
    red_left_cm: float, red_top_cm: float,
    red_right_cm: float, red_bottom_cm: float,
    blue_x_cm: float, page_height_cm: float, page_width_cm: float,
    green_top_y_cm: float, green_bottom_y_cm: float
) -> None:

    # Add debug visualization shapes to a header.

    # Clear existing header content.

    for paragraph in header.paragraphs:
        p = paragraph._p
        p.getparent ().remove ( p )

    # Add a paragraph for the drawing with minimal height.

    paragraph = header.add_paragraph ()

    # Minimize paragraph height so the header doesn't push content down.

    paragraph.paragraph_format.space_before      = Pt ( 0 )
    paragraph.paragraph_format.space_after       = Pt ( 0 )
    paragraph.paragraph_format.line_spacing      = Pt ( 1 )
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    run           = paragraph.add_run ()
    run.font.size = Pt ( 1 )

    # Convert cm to inches for VML (1 inch = 2.54 cm).

    def cm_to_in ( cm ):

        return cm / 2.54

    # Red rectangle dimensions.

    red_width_in  = cm_to_in ( red_right_cm - red_left_cm )
    red_height_in = cm_to_in ( red_bottom_cm - red_top_cm )
    red_left_in   = cm_to_in ( red_left_cm )
    red_top_in    = cm_to_in ( red_top_cm )

    # Blue line position.

    blue_x_in      = cm_to_in ( blue_x_cm )
    page_height_in = cm_to_in ( page_height_cm )

    # Green line positions.

    page_width_in     = cm_to_in ( page_width_cm )
    green_top_y_in    = cm_to_in ( green_top_y_cm )
    green_bottom_y_in = cm_to_in ( green_bottom_y_cm )

    # Convert to EMUs for DrawingML (1 inch = 914400 EMUs).

    EMU_PER_IN         = 914400
    red_left_emu       = int ( red_left_in * EMU_PER_IN )
    red_top_emu        = int ( red_top_in * EMU_PER_IN )
    red_width_emu      = int ( red_width_in * EMU_PER_IN )
    red_height_emu     = int ( red_height_in * EMU_PER_IN )
    blue_x_emu         = int ( blue_x_in * EMU_PER_IN )
    page_height_emu    = int ( page_height_in * EMU_PER_IN )
    page_width_emu     = int ( page_width_in * EMU_PER_IN )
    green_top_y_emu    = int ( green_top_y_in * EMU_PER_IN )
    green_bottom_y_emu = int ( green_bottom_y_in * EMU_PER_IN )

    # Create DrawingML shapes with explicit "in front of text" positioning.
    # relativeHeight controls z-order (higher = more in front).
    # behindDoc="0" means in front of document content.
    # Z-order (back to front): green lines -> blue line -> red rectangle.

    # Red rectangle (highest z-order to appear on top).

    red_rect_xml = f'''
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
               xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
               xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="251659268" behindDoc="0" locked="0"
                   layoutInCell="0" allowOverlap="1"
                   wp14:anchorId="1A000000" wp14:editId="1A000001">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>{red_left_emu}</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>{red_top_emu}</wp:posOffset></wp:positionV>
            <wp:extent cx="{red_width_emu}" cy="{red_height_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="1" name="DebugRedRect"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                    <wps:wsp>
                        <wps:cNvSpPr/>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="{red_width_emu}" cy="{red_height_emu}"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                            <a:noFill/>
                            <a:ln w="6350">
                                <a:solidFill><a:srgbClr val="FF0000"/></a:solidFill>
                            </a:ln>
                        </wps:spPr>
                        <wps:bodyPr/>
                    </wps:wsp>
                </a:graphicData>
            </a:graphic>
        </wp:anchor>
    </w:drawing>
    '''

    # Blue vertical line.

    blue_line_xml = f'''
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
               xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="251659265" behindDoc="0" locked="0"
                   layoutInCell="0" allowOverlap="1"
                   wp14:anchorId="1A000002" wp14:editId="1A000003">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>{blue_x_emu}</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
            <wp:extent cx="0" cy="{page_height_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="2" name="DebugBlueLine"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                    <wps:wsp>
                        <wps:cNvCnPr/>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="0" cy="{page_height_emu}"/>
                            </a:xfrm>
                            <a:prstGeom prst="line"><a:avLst/></a:prstGeom>
                            <a:ln w="6350">
                                <a:solidFill><a:srgbClr val="0000FF"/></a:solidFill>
                            </a:ln>
                        </wps:spPr>
                        <wps:bodyPr/>
                    </wps:wsp>
                </a:graphicData>
            </a:graphic>
        </wp:anchor>
    </w:drawing>
    '''

    # Green horizontal line (top) - header text boundary.

    green_top_line_xml = f'''
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
               xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="251659266" behindDoc="0" locked="0"
                   layoutInCell="0" allowOverlap="1"
                   wp14:anchorId="1A000004" wp14:editId="1A000005">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>{green_top_y_emu}</wp:posOffset></wp:positionV>
            <wp:extent cx="{page_width_emu}" cy="0"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="3" name="DebugGreenTopLine"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                    <wps:wsp>
                        <wps:cNvCnPr/>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="{page_width_emu}" cy="0"/>
                            </a:xfrm>
                            <a:prstGeom prst="line"><a:avLst/></a:prstGeom>
                            <a:ln w="6350">
                                <a:solidFill><a:srgbClr val="00FF00"/></a:solidFill>
                            </a:ln>
                        </wps:spPr>
                        <wps:bodyPr/>
                    </wps:wsp>
                </a:graphicData>
            </a:graphic>
        </wp:anchor>
    </w:drawing>
    '''

    # Green horizontal line (bottom) - footer text boundary.

    green_bottom_line_xml = f'''
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
               xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="251659267" behindDoc="0" locked="0"
                   layoutInCell="0" allowOverlap="1"
                   wp14:anchorId="1A000006" wp14:editId="1A000007">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>{green_bottom_y_emu}</wp:posOffset></wp:positionV>
            <wp:extent cx="{page_width_emu}" cy="0"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="4" name="DebugGreenBottomLine"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                    <wps:wsp>
                        <wps:cNvCnPr/>
                        <wps:spPr>
                            <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="{page_width_emu}" cy="0"/>
                            </a:xfrm>
                            <a:prstGeom prst="line"><a:avLst/></a:prstGeom>
                            <a:ln w="6350">
                                <a:solidFill><a:srgbClr val="00FF00"/></a:solidFill>
                            </a:ln>
                        </wps:spPr>
                        <wps:bodyPr/>
                    </wps:wsp>
                </a:graphicData>
            </a:graphic>
        </wp:anchor>
    </w:drawing>
    '''

    # Append the drawing shapes to the paragraph runs.

    red_drawing = parse_xml ( red_rect_xml )
    run._r.append ( red_drawing )

    run2           = paragraph.add_run ()
    run2.font.size = Pt ( 1 )
    blue_drawing   = parse_xml ( blue_line_xml )
    run2._r.append ( blue_drawing )

    run3              = paragraph.add_run ()
    run3.font.size    = Pt ( 1 )
    green_top_drawing = parse_xml ( green_top_line_xml )
    run3._r.append ( green_top_drawing )

    run4                 = paragraph.add_run ()
    run4.font.size       = Pt ( 1 )
    green_bottom_drawing = parse_xml ( green_bottom_line_xml )
    run4._r.append ( green_bottom_drawing )


#-----------------------------------------------------------------------------------------------------------------------
# Function: add_config_info_overlay
#
# Description:
#
#   Add a configuration info text box to the current page in the document body.
#
#   Inserts a text box directly into the document (not the header) so it appears in front of all other content including
#   tables and images. Call this function on pages where you want the overlay to appear.
#
# Arguments:
#
#   document         : The document to add the text box to.
#   config           : Configuration with all settings to display.
#   is_recto         : True for recto (odd/right) pages, False for verso (even/left) pages. Controls whether the overlay
#                      appears bottom-right or bottom-left.
#   anchor_paragraph : Optional existing paragraph to anchor the text box to. Use this when the page is already full
#                      (e.g., full-page images) to avoid creating a new paragraph that would go to the next page.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def add_config_info_overlay ( document: Document, config: Config, is_recto: bool = True,
                              anchor_paragraph = None ) -> None:

    # Add a configuration info text box to the current page in the document body.

    if not config.debug.config_info_overlay:
        return

    # Get measurements.

    overlay_bottom_cm = config.config_info_overlay.bottom
    overlay_right_cm  = config.config_info_overlay.right
    overlay_left_cm   = config.config_info_overlay.left
    overlay_title     = config.config_info_overlay.title
    page_width_cm     = config.page.width
    page_height_cm    = config.page.height

    # Build config info text.

    config_text = _build_config_info_text ( config )

    # Get settings from config.

    title_font_size   = config.config_info_overlay.title_font_size
    data_font_size    = config.config_info_overlay.data_font_size
    text_box_width_cm = config.config_info_overlay.width

    # Calculate text box dimensions based on content.

    lines      = [ overlay_title, '' ] + config_text.split ( '\n' )
    line_count = len ( lines )

    # Calculate exact height based on line heights (matching the exact lineRule in make_para).
    # Line heights: title uses title_font_size * 1.2, fields use data_font_size * 1.2.
    # Convert points to cm: points / 72 * 2.54.

    title_line_height_cm = title_font_size * 1.2 / 72 * 2.54
    field_line_height_cm = data_font_size * 1.2 / 72 * 2.54
    text_height_cm       = title_line_height_cm + field_line_height_cm * ( line_count - 1 )

    # Text box internal margins (from bodyPr: tIns="45720" bIns="45720" EMUs).
    # 914400 EMU = 1 inch = 2.54 cm, so 45720 EMU = 0.127 cm.

    vertical_padding_cm = 2 * ( 45720 / 914400 * 2.54 )  # Top + bottom insets
    estimated_height_cm = text_height_cm + vertical_padding_cm

    # Use the configured width for the text box.

    estimated_width_cm = text_box_width_cm

    # Position based on page side.
    # Calculate x position for bottom-right (recto) or bottom-left (verso).

    if is_recto:

        # Recto (odd) pages: bottom-right.
        # Place right edge of text box at overlay_right from physical page right.

        x_cm = page_width_cm - overlay_right_cm - estimated_width_cm
    else:

        # Verso (even) pages: bottom-left.
        # Place left edge of text box at overlay_left from physical page left.

        x_cm = overlay_left_cm

    # Y position: distance from bottom of page.

    y_cm = page_height_cm - overlay_bottom_cm - estimated_height_cm

    # Add the text box to the document body.

    _add_config_textbox_to_body (
        document, x_cm, y_cm,
        estimated_width_cm, estimated_height_cm,
        overlay_title, config_text,
        title_font_size, data_font_size,
        anchor_paragraph = anchor_paragraph
    )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _build_config_info_text
#
# Description:
#
#   Build the text content for the config info overlay.
#
#   Lists all configuration values in a readable format.
#
# Arguments:
#
#   config : Configuration object with all settings.
#
# Returns:
#
#   Formatted string with all config values.
#-----------------------------------------------------------------------------------------------------------------------

def _build_config_info_text ( config: Config ) -> str:

    # Build the text content for the config info overlay.

    lines = []

    # Document info.

    lines.append ( f"document.title: {config.document.title}" )
    lines.append ( f"document.version: {config.document.version}" )
    lines.append ( f"document.year: {config.document.year}" )
    lines.append ( "" )

    # Page layout.

    lines.append ( f"page.width: {config.page.width} cm" )
    lines.append ( f"page.height: {config.page.height} cm" )
    lines.append ( f"page.margin_top: {config.page.margin_top} cm" )
    lines.append ( f"page.margin_bottom: {config.page.margin_bottom} cm" )
    lines.append ( f"page.margin_left: {config.page.margin_left} cm" )
    lines.append ( f"page.margin_right: {config.page.margin_right} cm" )
    lines.append ( f"page.gutter_size: {config.page.gutter_size} cm" )
    lines.append ( f"page.page_number_position: {config.page.page_number_position} cm" )
    lines.append ( "" )

    # Table styling - border.

    lines.append ( f"table.border.thickness: {config.table.border.thickness} pt" )
    lines.append ( f"table.border.grayscale: {config.table.border.grayscale}" )
    lines.append ( "" )

    # Table styling - title row.

    lines.append ( f"table.title_row.height: {config.table.title_row.height} pt" )
    lines.append ( f"table.title_row.background_grayscale: {config.table.title_row.background_grayscale}" )
    lines.append ( f"table.title_row.font_size: {config.table.title_row.font_size} pt" )
    lines.append ( f"table.title_row.font_grayscale: {config.table.title_row.font_grayscale}" )
    lines.append ( "" )

    # Table styling - header row.

    lines.append ( f"table.header_row.height: {config.table.header_row.height} pt" )
    lines.append ( f"table.header_row.background_grayscale: {config.table.header_row.background_grayscale}" )
    lines.append ( f"table.header_row.font_size: {config.table.header_row.font_size} pt" )
    lines.append ( f"table.header_row.font_grayscale: {config.table.header_row.font_grayscale}" )
    lines.append ( "" )

    # Table styling - content row.

    lines.append ( f"table.content_row.font_size: {config.table.content_row.font_size} pt" )
    lines.append ( f"table.content_row.font_grayscale: {config.table.content_row.font_grayscale}" )
    lines.append ( f"table.content_row.font_italic: {config.table.content_row.font_italic}" )
    lines.append ( "" )

    # Config info overlay.

    lines.append ( f"config_info_overlay.bottom: {config.config_info_overlay.bottom} cm" )
    lines.append ( f"config_info_overlay.right: {config.config_info_overlay.right} cm" )
    lines.append ( f"config_info_overlay.left: {config.config_info_overlay.left} cm" )
    lines.append ( f"config_info_overlay.width: {config.config_info_overlay.width} cm" )
    lines.append ( f"config_info_overlay.title_font_size: {config.config_info_overlay.title_font_size} pt" )
    lines.append ( f"config_info_overlay.data_font_size: {config.config_info_overlay.data_font_size} pt" )
    lines.append ( "" )

    # Section configs from raw.

    raw = config.raw

    # Cover contact table.

    if 'cover' in raw and 'contact_table' in raw [ 'cover' ]:

        ct = raw [ 'cover' ] [ 'contact_table' ]
        lines.append ( f"cover.contact_table.label_grayscale: {ct.get('label_grayscale', 'N/A')}" )
        lines.append ( "" )

    # TOC.

    if 'toc' in raw:

        toc = raw [ 'toc' ]
        lines.append ( f"toc.rows_per_page: {toc.get('rows_per_page', 'N/A')}" )
        lines.append ( f"toc.section_grayscale: {toc.get('section_grayscale', 'N/A')}" )
        lines.append ( f"toc.first_item_grayscale: {toc.get('first_item_grayscale', 'N/A')}" )
        lines.append ( "" )

    # Calendar.

    if 'calendar' in raw:

        cal = raw [ 'calendar' ]

        lines.append ( f"calendar.day_row_height: {cal.get('day_row_height', 'N/A')} pt" )
        lines.append ( f"calendar.month_name_gap: {cal.get('month_name_gap', 'N/A')} pt" )
        lines.append ( "" )

    # Week planner.

    if 'week_planner' in raw:

        wp = raw [ 'week_planner' ]

        lines.append ( f"week_planner.rows_per_page: {wp.get('rows_per_page', 'N/A')}" )
        lines.append ( f"week_planner.first_week_grayscale: {wp.get('first_week_grayscale', 'N/A')}" )
        lines.append ( "" )

    # Goals.

    if 'goals' in raw:

        g = raw [ 'goals' ]

        lines.append ( f"goals.columns: {g.get('columns', 'N/A')}" )
        lines.append ( f"goals.rows: {g.get('rows', 'N/A')}" )
        lines.append ( "" )

    # Backlog.

    if 'backlog' in raw:

        bl = raw [ 'backlog' ]

        lines.append ( f"backlog.page_count: {bl.get('page_count', 'N/A')}" )
        lines.append ( f"backlog.row_count: {bl.get('row_count', 'N/A')}" )
        lines.append ( "" )

    # Daily spread.

    if 'daily_spread' in raw:

        ds = raw [ 'daily_spread' ]

        lines.append ( f"daily_spread.rows: {ds.get('rows', 'N/A')}" )
        lines.append ( f"daily_spread.subject_width_percent: {ds.get('subject_width_percent', 'N/A')}%" )
        lines.append ( f"daily_spread.table_gap: {ds.get('table_gap', 'N/A')} cm" )
        lines.append ( "" )

    # Graph paper.

    if 'graph_paper' in raw:

        gp = raw [ 'graph_paper' ]

        lines.append ( f"graph_paper.page_count: {gp.get('page_count', 'N/A')}" )
        lines.append ( f"graph_paper.columns: {gp.get('columns', 'N/A')}" )
        lines.append ( f"graph_paper.rows: {gp.get('rows', 'N/A')}" )
        lines.append ( f"graph_paper.grid_color_percent: {gp.get('grid_color_percent', 'N/A')}" )
        lines.append ( f"graph_paper.border_color_percent: {gp.get('border_color_percent', 'N/A')}" )

    # Return data to caller.

    return '\n'.join ( lines )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_config_textbox_to_body
#
# Description:
#
#   Add a config info text box to the document body.
#
#   Inserts a floating text box positioned absolutely on the page, set to "in front of text" so it appears on top of all
#   content.
#
# Arguments:
#
#   document         : The document to add the text box to.
#   x_cm             : X position from page left (cm).
#   y_cm             : Y position from page top (cm).
#   width_cm         : Text box width (cm).
#   height_cm        : Text box height (cm).
#   title            : Title text for the overlay.
#   content          : Config info content text.
#   title_font_size  : Font size for title in points.
#   data_font_size   : Font size for field data in points.
#   anchor_paragraph : Optional existing paragraph to anchor the text box to. If None, creates a new minimal paragraph.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_config_textbox_to_body (
    document: Document,
    x_cm: float, y_cm: float,
    width_cm: float, height_cm: float,
    title: str, content: str,
    title_font_size: float, data_font_size: float,
    anchor_paragraph = None
) -> None:

    # Add a config info text box to the document body.

    import random

    if anchor_paragraph is not None:

        # Use the existing paragraph - add a run to it for the text box.

        paragraph     = anchor_paragraph
        run           = paragraph.add_run ()
        run.font.size = Pt ( 1 )
    else:

        # Create a new minimal paragraph to hold the text box.

        paragraph                                    = document.add_paragraph ()
        paragraph.paragraph_format.space_before      = Pt ( 0 )
        paragraph.paragraph_format.space_after       = Pt ( 0 )
        paragraph.paragraph_format.line_spacing      = Pt ( 1 )
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        run           = paragraph.add_run ()
        run.font.size = Pt ( 1 )

    # Convert cm to EMUs (914400 EMUs per inch, 1 inch = 2.54 cm).

    EMU_PER_CM = 914400 / 2.54
    x_emu      = int ( x_cm * EMU_PER_CM )
    y_emu      = int ( y_cm * EMU_PER_CM )
    width_emu  = int ( width_cm * EMU_PER_CM )
    height_emu = int ( height_cm * EMU_PER_CM )

    # Unique ID for this text box.

    doc_pr_id = random.randint ( 10000, 99999 )

    # Build text content - font sizes in half-points (Word uses half-points internally).

    title_size_half_pt = int ( title_font_size * 2 )
    field_size_half_pt = int ( data_font_size * 2 )
    font_name          = "Arial"

    # Escape XML special characters in text.

    def escape_xml ( text: str ) -> str:

        return text.replace ( '&', '&amp;' ).replace ( '<', '&lt;' ).replace ( '>', '&gt;' )

    # Create a run element with optional bold formatting and specified size.

    def make_run ( text: str, bold: bool = False, size_half_pt: int = 10 ) -> str:

        escaped  = escape_xml ( text )
        bold_tag = '<w:b/>' if bold else ''
        return f'''<w:r><w:rPr><w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/><w:sz w:val="{size_half_pt}"/><w:szCs w:val="{size_half_pt}"/>{bold_tag}</w:rPr><w:t xml:space="preserve">{escaped}</w:t></w:r>'''

    # Calculate line heights in twips (20 twips per point, with 1.2x line spacing).

    title_line_height_twips = int ( title_font_size * 1.2 * 20 )
    field_line_height_twips = int ( data_font_size * 1.2 * 20 )

    # Create a paragraph. Title uses larger font, fields use smaller font.

    def make_para ( text: str, is_title: bool = False ) -> str:

        if is_title:

            # Title line - larger font, bold, exact line height.

            return f'''<w:p><w:pPr><w:spacing w:after="0" w:line="{title_line_height_twips}" w:lineRule="exact"/></w:pPr>{make_run(text, bold=True, size_half_pt=title_size_half_pt)}</w:p>'''
        elif ':' not in text:

            # Empty line or line without colon - exact line height matching data fields.

            return f'''<w:p><w:pPr><w:spacing w:after="0" w:line="{field_line_height_twips}" w:lineRule="exact"/></w:pPr>{make_run(text, bold=False, size_half_pt=field_size_half_pt)}</w:p>'''
        else:

            # Split at first colon - field name is bold, value is normal.
            # Add two spaces after colon for readability.

            colon_idx   = text.index ( ':' )
            field_name  = text [ : colon_idx + 1 ]  # Include the colon
            field_value = "  " + text [ colon_idx + 1 : ].lstrip ()  # Two spaces + value
            return f'''<w:p><w:pPr><w:spacing w:after="0" w:line="{field_line_height_twips}" w:lineRule="exact"/></w:pPr>{make_run(field_name, bold=True, size_half_pt=field_size_half_pt)}{make_run(field_value, bold=False, size_half_pt=field_size_half_pt)}</w:p>'''

    # Build content paragraphs (title + blank line + config values).

    lines       = [ title, '' ] + content.split ( '\n' )
    content_xml = ''.join ( make_para ( line, is_title = ( i == 0 ) ) for i, line in enumerate ( lines ) )

    # Standard Word text box XML - same as Insert > Text Box.
    # behindDoc="0" = in front of text.
    # relativeFrom="page" = positioned relative to page edges.

    textbox_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
        xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
      <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
                 simplePos="0" relativeHeight="251659264" behindDoc="0"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{doc_pr_id}" name="ConfigInfo_{doc_pr_id}"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks/></wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBox="1"><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
              <wps:spPr bwMode="auto">
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill>
                <a:ln w="6350"><a:solidFill><a:srgbClr val="000000"/></a:solidFill><a:miter lim="800000"/><a:headEnd/><a:tailEnd/></a:ln>
              </wps:spPr>
              <wps:txbx><w:txbxContent>{content_xml}</w:txbxContent></wps:txbx>
              <wps:bodyPr rot="0" vert="horz" wrap="square" lIns="91440" tIns="45720" rIns="91440" bIns="45720" anchor="t" anchorCtr="0" upright="1"><a:spAutoFit/></wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>'''

    # Append the text box drawing to the run.

    drawing = parse_xml ( textbox_xml )
    run._r.append ( drawing )


#-----------------------------------------------------------------------------------------------------------------------
# Function: add_section_break
#
# Description:
#
#   Add a section break and configure the new section with gutter.
#
#   Use this after the cover page to enable gutters for the main content. The new section inherits mirror margins from
#   the document settings.
#
# Arguments:
#
#   document : The document to add a section break to.
#   config   : Configuration with page settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def add_section_break ( document: Document, config: Config ) -> None:

    # Add a section break and configure the new section with gutter.

    from docx.enum.section import WD_SECTION

    # Add a new section (continuous section break, then we configure it).

    new_section = document.add_section ( WD_SECTION.NEW_PAGE )

    # Configure the new section with gutter enabled.

    configure_section ( new_section, config )


#-----------------------------------------------------------------------------------------------------------------------
# Function: add_non_numbered_section_break
#
# Description:
#
#   Add a section break without page numbers.
#
#   Creates a new section with empty footers (no page numbers). Use this for sections like the rear cover that should not
#   have page numbers.
#
# Arguments:
#
#   document : The document to add a section break to.
#   config   : Configuration with page settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def add_non_numbered_section_break ( document: Document, config: Config ) -> None:

    # Add a section break without page numbers.

    from docx.enum.section import WD_SECTION

    # Add a new section.

    new_section = document.add_section ( WD_SECTION.NEW_PAGE )

    # Configure the section with page layout.

    configure_section ( new_section, config )

    # Unlink footers from the previous section and clear them.
    # This ensures no page numbers appear in this section.

    odd_footer                       = new_section.footer
    odd_footer.is_linked_to_previous = False

    # Clear any existing content.

    for paragraph in odd_footer.paragraphs:
        paragraph.clear ()

    even_footer                       = new_section.even_page_footer
    even_footer.is_linked_to_previous = False

    # Clear any existing content.

    for paragraph in even_footer.paragraphs:
        paragraph.clear ()


#-----------------------------------------------------------------------------------------------------------------------
# Function: add_numbered_section_break
#
# Description:
#
#   Add a section break with page numbering enabled.
#
#   Creates a new section with:
#   - Different odd/even footers for recto/verso page number positioning.
#   - Recto (odd) pages: page number bottom-right, right-aligned.
#   - Verso (even) pages: page number bottom-left, left-aligned.
#   - Page numbering starting at specified number.
#
# Arguments:
#
#   document     : The document to add a section break to.
#   config       : Configuration with page settings.
#   start_number : Starting page number (default: 1).
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def add_numbered_section_break ( document: Document, config: Config, start_number: int = 1 ) -> None:

    # Add a section break with page numbering enabled.

    from docx.enum.section import WD_SECTION
    from docx.enum.text    import WD_PARAGRAPH_ALIGNMENT

    # Add a new section.

    new_section = document.add_section ( WD_SECTION.NEW_PAGE )

    # Configure the section with page layout.

    configure_section ( new_section, config )

    # Enable different odd and even headers/footers.

    _enable_different_odd_even_headers ( new_section )

    # Set the starting page number.

    _set_page_number_start ( new_section, start_number )

    # Configure the odd (recto) footer - page number right-aligned.

    odd_footer                       = new_section.footer
    odd_footer.is_linked_to_previous = False
    _add_page_number_to_footer ( odd_footer, WD_PARAGRAPH_ALIGNMENT.RIGHT, config )

    # Configure the even (verso) footer - page number left-aligned.

    even_footer                       = new_section.even_page_footer
    even_footer.is_linked_to_previous = False
    _add_page_number_to_footer ( even_footer, WD_PARAGRAPH_ALIGNMENT.LEFT, config )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _enable_different_odd_even_headers
#
# Description:
#
#   Enable different odd and even headers/footers for a section.
#
#   This setting affects the entire document but is set per-section. Word will then use odd_page_header/footer and
#   even_page_header/footer.
#
# Arguments:
#
#   section : The section to configure.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _enable_different_odd_even_headers ( section: Section ) -> None:

    # Enable different odd and even headers/footers for a section.

    # Access the sectPr element.

    sectPr = section._sectPr

    # Find or create the evenAndOddHeaders element in document settings.
    # This is actually a document-level setting, but we access it here.

    document = section._document_part.document
    settings = document.settings.element

    # Check if evenAndOddHeaders already exists.

    evenAndOdd = settings.find ( qn ( 'w:evenAndOddHeaders' ) )

    if evenAndOdd is None:

        evenAndOdd = parse_xml (
            '<w:evenAndOddHeaders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )

        settings.append ( evenAndOdd )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _set_page_number_start
#
# Description:
#
#   Set the starting page number for a section.
#
# Arguments:
#
#   section      : The section to configure.
#   start_number : The page number to start from.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _set_page_number_start ( section: Section, start_number: int ) -> None:

    # Set the starting page number for a section.

    sectPr = section._sectPr

    # Find or create the pgNumType element.

    pgNumType = sectPr.find ( qn ( 'w:pgNumType' ) )

    if pgNumType is None:

        pgNumType = parse_xml (
            f'<w:pgNumType xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:start="{start_number}"/>'
        )

        sectPr.append ( pgNumType )

    else:

        pgNumType.set ( qn ( 'w:start' ), str ( start_number ) )


#-----------------------------------------------------------------------------------------------------------------------
# Function: _add_page_number_to_footer
#
# Description:
#
#   Add a page number field to a footer.
#
# Arguments:
#
#   footer    : The footer to add the page number to.
#   alignment : WD_PARAGRAPH_ALIGNMENT value for text alignment.
#   config    : Configuration with page settings.
#
# Returns:
#
#   None.
#-----------------------------------------------------------------------------------------------------------------------

def _add_page_number_to_footer ( footer, alignment, config: Config ) -> None:

    # Add a page number field to a footer.

    # Clear any existing content.

    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent ().remove ( p )

    # Add the paragraph with page number.

    paragraph           = footer.add_paragraph ()
    paragraph.alignment = alignment

    # Set paragraph format - no spacing.

    paragraph.paragraph_format.space_before = Pt ( 0 )
    paragraph.paragraph_format.space_after  = Pt ( 0 )

    # Add the PAGE field for page number.

    run           = paragraph.add_run ()
    run.font.name = DEFAULT_FONT_NAME
    run.font.size = Pt ( 10 )

    # Insert the PAGE field using XML.
    # The PAGE field displays the current page number.

    fldChar_begin = parse_xml (
        '<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:fldCharType="begin"/>'
    )
    run._r.append ( fldChar_begin )

    instrText = parse_xml (
        '<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xml:space="preserve"> PAGE </w:instrText>'
    )
    run._r.append ( instrText )

    fldChar_end = parse_xml (
        '<w:fldChar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:fldCharType="end"/>'
    )
    run._r.append ( fldChar_end )


#-----------------------------------------------------------------------------------------------------------------------
# Function: add_page_break
#
# Description:
#
#   Add a page break to the document.
#
# Arguments:
#
#   document        : The document to add a page break to.
#   minimize_height : If True, minimize the paragraph height to near-zero. Use this when precise page filling is needed.
#
# Returns:
#
#   The paragraph containing the page break (can be used as anchor for overlays).
#-----------------------------------------------------------------------------------------------------------------------

def add_page_break ( document: Document, minimize_height: bool = False ):

    # Add a page break to the document.

    from docx.shared    import Pt
    from docx.enum.text import WD_BREAK, WD_LINE_SPACING

    paragraph = document.add_paragraph ()
    run       = paragraph.add_run ()
    run.add_break ( WD_BREAK.PAGE )  # Page break (not line break)

    if minimize_height:

        # Set font size to 1pt and line spacing to exactly 1pt.
        # Also set spacing before/after to 0 to eliminate any extra space.

        run.font.size                                = Pt ( 1 )
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing      = Pt ( 1 )
        paragraph.paragraph_format.space_before      = Pt ( 0 )
        paragraph.paragraph_format.space_after       = Pt ( 0 )

    # Return data to caller.

    return paragraph


#-----------------------------------------------------------------------------------------------------------------------
# Function: get_content_width
#
# Description:
#
#   Calculate the available content width in centimeters.
#
#   By default, includes gutter in calculation since tables and content must fit within the actual printable area (page
#   width minus margins minus gutter on binding side).
#
# Arguments:
#
#   config         : Configuration with page settings.
#   include_gutter : Whether to subtract gutter from width (default: True).
#
# Returns:
#
#   Content width in centimeters.
#-----------------------------------------------------------------------------------------------------------------------

def get_content_width ( config: Config, include_gutter: bool = True ) -> float:

    # Calculate the available content width in centimeters.

    width = config.page.width - config.page.margin_left - config.page.margin_right

    if include_gutter:

        width -= config.page.gutter_size

    # Return data to caller.

    return width


#-----------------------------------------------------------------------------------------------------------------------
# Function: get_content_width_twips
#
# Description:
#
#   Calculate the available content width in twips (dxa).
#
#   Always includes gutter since this is primarily used for table column width calculations where content must fit
#   within the printable area.
#
# Arguments:
#
#   config : Configuration with page settings.
#
# Returns:
#
#   Content width in twips (1 cm ≈ 566.93 twips).
#-----------------------------------------------------------------------------------------------------------------------

def get_content_width_twips ( config: Config ) -> int:

    # Calculate the available content width in twips (dxa).

    width_cm = get_content_width ( config, include_gutter = True )

    # Return data to caller.

    return int ( width_cm * TWIPS_PER_CM )


#-----------------------------------------------------------------------------------------------------------------------
# Function: get_content_height
#
# Description:
#
#   Calculate the available content height in centimeters.
#
# Arguments:
#
#   config : Configuration with page settings.
#
# Returns:
#
#   Content height in centimeters.
#-----------------------------------------------------------------------------------------------------------------------

def get_content_height ( config: Config ) -> float:

    # Calculate the available content height in centimeters.

    # Return data to caller.

    return config.page.height - config.page.margin_top - config.page.margin_bottom


#-----------------------------------------------------------------------------------------------------------------------
# Function: get_content_height_twips
#
# Description:
#
#   Calculate the available content height in twips.
#
# Arguments:
#
#   config : Configuration with page settings.
#
# Returns:
#
#   Content height in twips (1 cm ≈ 566.93 twips).
#-----------------------------------------------------------------------------------------------------------------------

def get_content_height_twips ( config: Config ) -> float:

    # Calculate the available content height in twips.

    height_cm = config.page.height - config.page.margin_top - config.page.margin_bottom

    # Return data to caller.

    return height_cm * TWIPS_PER_CM

# Global constants.

# Height of an empty paragraph with default font (Times New Roman 11pt).
# Word's "Single" line spacing is ~1.08-1.15x font size, so ~12-13pt for 11pt font.

EMPTY_PARAGRAPH_HEIGHT_TWIPS = 240  # ~12pt conservative estimate

# Height of a minimized paragraph (1pt font with exact 1pt line spacing).

MINIMIZED_PARAGRAPH_HEIGHT_TWIPS = 20  # 1pt = 20 twips

# Safety margin to prevent overflow due to rounding or unexpected spacing.
# This ensures tables never spill to the next page.

SAFETY_MARGIN_TWIPS = 40  # 2pt buffer


#-----------------------------------------------------------------------------------------------------------------------
# Function: get_title_row_height_twips
#
# Description:
#
#   Get title row height in twips from config.
#
# Arguments:
#
#   config : Configuration with table settings.
#
# Returns:
#
#   Title row height in twips.
#-----------------------------------------------------------------------------------------------------------------------

def get_title_row_height_twips ( config: Config ) -> int:

    # Get title row height in twips from config.

    # Return data to caller.

    return int ( config.table.title_row.height * TWIPS_PER_PT )


#-----------------------------------------------------------------------------------------------------------------------
# Function: get_header_row_height_twips
#
# Description:
#
#   Get header row height in twips from config.
#
# Arguments:
#
#   config : Configuration with table settings.
#
# Returns:
#
#   Header row height in twips.
#-----------------------------------------------------------------------------------------------------------------------

def get_header_row_height_twips ( config: Config ) -> int:

    # Get header row height in twips from config.

    # Return data to caller.

    return int ( config.table.header_row.height * TWIPS_PER_PT )


#-----------------------------------------------------------------------------------------------------------------------
# Function: grayscale_to_hex
#
# Description:
#
#   Convert grayscale percentage to hex color string.
#
# Arguments:
#
#   grayscale : Grayscale value (0=white, 100=black).
#
# Returns:
#
#   Hex color string (e.g., "000000" for black, "FFFFFF" for white).
#-----------------------------------------------------------------------------------------------------------------------

def grayscale_to_hex ( grayscale: int ) -> str:

    # Convert grayscale percentage to hex color string.

    gray_value = int ( 255 * ( 1 - grayscale / 100 ) )

    # Return data to caller.

    return f"{gray_value:02X}{gray_value:02X}{gray_value:02X}"


#-----------------------------------------------------------------------------------------------------------------------
# Function: grayscale_to_rgb
#
# Description:
#
#   Convert grayscale percentage to RGB tuple.
#
# Arguments:
#
#   grayscale : Grayscale value (0=white, 100=black).
#
# Returns:
#
#   RGB tuple (r, g, b) where each value is 0-255.
#-----------------------------------------------------------------------------------------------------------------------

def grayscale_to_rgb ( grayscale: int ) -> tuple [ int, int, int ]:

    # Convert grayscale percentage to RGB tuple.

    gray_value = int ( 255 * ( 1 - grayscale / 100 ) )

    # Return data to caller.

    return ( gray_value, gray_value, gray_value )


#-----------------------------------------------------------------------------------------------------------------------
# Function: compute_table_row_height
#
# Description:
#
#   Compute the content row height so the table fills the page with a safety margin.
#
#   Uses the formula: p_v = p_para + r_t + r_h + r_c * n + safety
#   Solving for r_c: r_c = (p_v - p_para - r_t - r_h - safety) / n
#
# Arguments:
#
#   config                           : Configuration with page settings.
#   num_content_rows                 : Number of content rows (n).
#   title_row_height_twips           : Height of title row in twips (r_t).
#   header_row_height_twips          : Height of header row in twips (r_h).
#   preceding_paragraph_height_twips : Height of any preceding paragraph in twips. Use
#                                      MINIMIZED_PARAGRAPH_HEIGHT_TWIPS (20) for minimized page breaks, or
#                                      EMPTY_PARAGRAPH_HEIGHT_TWIPS (240) for normal paragraphs.
#
# Returns:
#
#   Computed content row height in twips (r_c).
#-----------------------------------------------------------------------------------------------------------------------

def compute_table_row_height (
    config: Config,
    num_content_rows: int,
    title_row_height_twips: int,
    header_row_height_twips: int,
    preceding_paragraph_height_twips: int = 0
) -> int:

    # Compute the content row height so the table fills the page with a safety margin.

    p_v = get_content_height_twips ( config )

    # Account for the preceding paragraph and safety margin.

    p_v -= preceding_paragraph_height_twips
    p_v -= SAFETY_MARGIN_TWIPS

    # Compute the content row height.

    available_for_content = p_v - title_row_height_twips - header_row_height_twips
    row_height            = available_for_content / num_content_rows

    # Return as integer (twips must be whole numbers).

    return int ( row_height )


#-----------------------------------------------------------------------------------------------------------------------
# Function: validate_table_height
#
# Description:
#
#   Validate and compute table row height, warning if content doesn't fit.
#
#   This function computes the content row height using the same formula as compute_table_row_height(), but also checks
#   if the result is valid and prints a warning to the terminal if the table won't fit on the page.
#
#   Use this function when you want to catch configuration issues early, particularly before adding page numbers (which
#   reduces available space).
#
# Arguments:
#
#   config                           : Configuration with page settings.
#   section_name                     : Name of the section for warning messages (e.g., "Week Planner").
#   num_content_rows                 : Number of content rows (n).
#   title_row_height_twips           : Height of title row in twips (r_t).
#   header_row_height_twips          : Height of header row in twips (r_h).
#   preceding_paragraph_height_twips : Height of any preceding paragraph in twips.
#
# Returns:
#
#   Computed content row height in twips. Returns the value even if invalid so callers can decide how to handle the
#   error.
#-----------------------------------------------------------------------------------------------------------------------

def validate_table_height (
    config: Config,
    section_name: str,
    num_content_rows: int,
    title_row_height_twips: int,
    header_row_height_twips: int,
    preceding_paragraph_height_twips: int = 0
) -> int:

    # Validate and compute table row height, warning if content doesn't fit.

    available = get_content_height_twips ( config )
    fixed_overhead = (
        preceding_paragraph_height_twips +
        SAFETY_MARGIN_TWIPS +
        title_row_height_twips +
        header_row_height_twips
    )
    remaining  = available - fixed_overhead
    row_height = int ( remaining / num_content_rows )

    # Warn if the computed row height is invalid (table won't fit on the page).

    if row_height <= 0:
        print ( f"\n(!) WARNING: {section_name} table height validation failed!" )
        print ( f"    Available space: {available:.0f} twips ({available / TWIPS_PER_CM:.2f} cm)" )
        print ( f"    Fixed overhead: {fixed_overhead:.0f} twips ({fixed_overhead / TWIPS_PER_CM:.2f} cm)" )
        print ( f"      - Preceding paragraph: {preceding_paragraph_height_twips} twips" )
        print ( f"      - Safety margin: {SAFETY_MARGIN_TWIPS} twips" )
        print ( f"      - Title row: {title_row_height_twips} twips" )
        print ( f"      - Header row: {header_row_height_twips} twips" )
        print ( f"    Remaining for {num_content_rows} content rows: {remaining:.0f} twips" )
        print ( f"    Computed row height: {row_height} twips (INVALID - must be > 0)" )
        print ( f"    Suggestion: Reduce margins, row counts, or row heights.\n" )

    # Return data to caller.

    return row_height
